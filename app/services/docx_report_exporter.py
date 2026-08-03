"""
翻译内容校对 Word 报告生成器

对标 app/services/xlsx_exporter.py 的结构，使用 python-docx（项目已有依赖）。

公开函数：
    build_translation_review_docx(report, items, runs) → bytes
    build_docx_download_response(filename, docx_bytes) → StreamingResponse
"""
from __future__ import annotations

from io import BytesIO
from urllib.parse import quote

from fastapi.responses import StreamingResponse

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_SEVERITY_LABEL = {"error": "错误", "warning": "警告", "suggestion": "建议"}
_STATUS_LABEL = {
    "open": "待处理", "applied": "已应用", "rejected": "已拒绝",
    "ignored": "已忽略", "stale": "需重查",
}
_CONF_LABEL = {"high": "高", "medium": "中", "low": "低"}


# ─── 公开入口 ──────────────────────────────────────────────

def build_translation_review_docx(
    report: dict,
    items: list[dict],
    runs: list[dict],
) -> bytes:
    """生成 Word 校对报告，返回字节。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── 标题 ────────────────────────────────────────────────
    _set_default_font(doc)
    title_para = doc.add_heading("翻译内容校对报告", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 概览 ────────────────────────────────────────────────
    doc.add_heading("一、概览", level=1)
    meta_rows = [
        ("范围", _scope_label(report)),
        ("检查范围", report.get("segment_scope") or "all"),
        ("句段总数", str(report.get("total_segments", 0))),
        ("问题总数", str(report.get("issue_count", 0))),
        ("待处理", str(report.get("active_issue_count", 0))),
        ("已应用", str(report.get("applied_count", 0))),
        ("已忽略", str(report.get("ignored_count", 0))),
        ("多类别句段", str(report.get("multi_category_segment_count", 0))),
        ("模型", report.get("model") or report.get("provider") or "—"),
        ("联网查证", report.get("web_verify_provider") or "none"),
        ("联网搜索次数", str(report.get("web_search_requests", 0))),
        ("生成时间", (report.get("created_at") or "")[:19].replace("T", " ")),
        ("完成时间", (report.get("finished_at") or "")[:19].replace("T", " ")),
    ]
    _two_col_table(doc, meta_rows)

    # ── 类别汇总 ─────────────────────────────────────────────
    doc.add_heading("二、类别汇总", level=1)
    cat_counts: dict = report.get("category_counts") or {}
    cat_headers = ["类别", "问题数", "检查状态"]
    cat_rows = []
    for run in sorted(runs, key=lambda r: r.get("category_index", 0)):
        key = run.get("category_key", "")
        label = run.get("label") or key
        count = cat_counts.get(key, 0)
        status_raw = run.get("status", "pending")
        status_display = {
            "ok": "✓ 完成", "skipped_no_candidate": "✓ 无需 AI",
            "partial": "⚠ 部分完成", "api_error": "✗ API 失败",
            "parse_failed": "✗ 解析失败", "skipped_no_prompt": "— 跳过",
            "pending": "—",
        }.get(status_raw, status_raw)
        cat_rows.append([label, str(count), status_display])
    _table(doc, cat_headers, cat_rows)

    # ── 按类别明细 ───────────────────────────────────────────
    doc.add_heading("三、按类别明细", level=1)
    from collections import defaultdict
    by_cat: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        by_cat[item.get("category_key", "")].append(item)

    detail_headers = ["句段", "文件", "规则", "原文", "译文（高亮问题片段）", "建议", "严重", "置信", "状态"]
    for run in sorted(runs, key=lambda r: r.get("category_index", 0)):
        key = run.get("category_key", "")
        label = run.get("label") or key
        cat_items = by_cat.get(key, [])
        if not cat_items:
            continue
        doc.add_heading(f"3.{run.get('category_index', 0)+1} {label}（{len(cat_items)} 条）", level=2)
        rows = []
        for it in cat_items:
            target_cell = _make_target_with_highlight(doc, it)
            seg_num = str(it.get("display_index", -1) + 1) if (it.get("display_index", -1) >= 0) else it.get("sentence_id", "")
            rows.append([
                seg_num,
                it.get("file_name") or "",
                it.get("rule_ref") or "",
                (it.get("source_text") or "")[:120],
                target_cell,
                (it.get("suggested_value") or it.get("suggested_target_text") or "")[:200],
                _SEVERITY_LABEL.get(it.get("severity", ""), it.get("severity", "")),
                _CONF_LABEL.get(it.get("confidence", ""), it.get("confidence", "")),
                _STATUS_LABEL.get(it.get("status", ""), it.get("status", "")),
            ])
        _table(doc, detail_headers, rows, narrow_first=True)

    # ── 按句段汇总（命中 ≥2 类别） ────────────────────────────
    multi_items = _collect_multi_category_segments(items)
    if multi_items:
        doc.add_heading("四、多类别问题句段汇总", level=1)
        doc.add_paragraph("以下句段命中 2 个或以上类别问题：")
        for sid_key, group in multi_items.items():
            first = group[0]
            seg_num = str(first.get("display_index", -1) + 1) if first.get("display_index", -1) >= 0 else first.get("sentence_id", "")
            p = doc.add_paragraph()
            run_obj = p.add_run(f"句段 {seg_num}（{first.get('file_name', '')}）")
            run_obj.bold = True
            doc.add_paragraph(f"  原文：{(first.get('source_text') or '')[:200]}", style="Quote")
            doc.add_paragraph(f"  译文：{(first.get('target_text') or '')[:200]}", style="Quote")
            for it in group:
                doc.add_paragraph(
                    f"  [{it.get('category_label', it.get('category_key', ''))}] §{it.get('rule_ref', '')} {it.get('reason', '')}",
                )

    # ── 附录 A：规则条款索引 ──────────────────────────────────
    doc.add_heading("附录 A：规则条款索引", level=1)
    for item in _RULE_INDEX:
        doc.add_paragraph(f"§{item[0]}  {item[1]}")

    # ── 附录 B：检查执行情况 ──────────────────────────────────
    doc.add_heading("附录 B：检查执行情况", level=1)
    runs_headers = ["类别", "输入句段", "AI 输入", "程序发现", "AI 发现", "丢弃", "状态"]
    runs_rows = [
        [
            r.get("label") or r.get("category_key", ""),
            str(r.get("input_segment_count", 0)),
            str(r.get("ai_input_count", 0)),
            str(r.get("program_finding_count", 0)),
            str(r.get("ai_finding_count", 0)),
            str(r.get("dropped_count", 0)),
            r.get("status", ""),
        ]
        for r in sorted(runs, key=lambda r: r.get("category_index", 0))
    ]
    _table(doc, runs_headers, runs_rows)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_docx_download_response(filename: str, docx_bytes: bytes) -> StreamingResponse:
    safe_filename = filename if filename.lower().endswith(".docx") else f"{filename}.docx"
    ascii_filename = safe_filename.encode("ascii", "ignore").decode("ascii").strip() or "report.docx"
    ascii_filename = ascii_filename.replace('"', "")
    quoted_filename = quote(safe_filename)
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type=DOCX_MEDIA_TYPE,
        headers={
            "Content-Disposition": (
                f'attachment; filename="{ascii_filename}"; '
                f"filename*=UTF-8''{quoted_filename}"
            )
        },
    )


# ─── 内部工具 ──────────────────────────────────────────────

def _set_default_font(doc) -> None:
    from docx.shared import Pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)


def _two_col_table(doc, rows: list[tuple[str, str]]) -> None:
    from docx.shared import Pt, RGBColor
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value
    doc.add_paragraph()


def _table(doc, headers: list[str], rows: list[list], narrow_first: bool = False) -> None:
    from docx.shared import Pt
    if not rows:
        doc.add_paragraph("（此类别无问题）")
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, cell_val in enumerate(row_data):
            if isinstance(cell_val, str):
                row[i].text = cell_val
    doc.add_paragraph()


def _make_target_with_highlight(doc, item: dict) -> str:
    """
    构造带高亮标注的译文字符串。
    由于 _table 目前只接受 str，这里先用简单标记；后续可升级为 Run 级 highlight。
    """
    target = item.get("target_text") or ""
    quote = item.get("quote") or ""
    locate = item.get("locate_status", "ok")
    if not quote or locate in ("unlocatable", "ambiguous"):
        return target[:300]
    start = item.get("quote_start", -1)
    end = item.get("quote_end", -1)
    if start >= 0 and end > start and end <= len(target):
        return target[:start] + f"【{target[start:end]}】" + target[end:300]
    # fallback: find by value
    idx = target.find(quote)
    if idx >= 0:
        return target[:idx] + f"【{quote}】" + target[idx + len(quote):300]
    return target[:300]


def _collect_multi_category_segments(items: list[dict]) -> dict[str, list[dict]]:
    """返回 {file_id:sentence_id: [items]} 且 len >= 2。"""
    from collections import defaultdict
    groups: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        key = f"{item.get('file_record_id', '')}:{item.get('sentence_id', '')}"
        groups[key].append(item)
    return {k: v for k, v in groups.items() if len(v) >= 2}


def _scope_label(report: dict) -> str:
    scope = report.get("scope", "file")
    if scope == "merge_view":
        n = report.get("total_files", 0)
        return f"合并视图（{n} 个文件）"
    return report.get("file_ids", [""])[0] if report.get("file_ids") else "单文件"


# ─── 规则索引（简版） ─────────────────────────────────────

_RULE_INDEX = [
    ("0.1", "时态：除客观事实描述外，一律用过去时或过去完成时"),
    ("0.2", "时态：图片说明与正文规则一致，用过去时"),
    ("1.1", "符号：不在英文译文中使用中文标点"),
    ("1.2", "符号：& 前后各空一格；/ 前后不空格"),
    ("1.3", "符号：电话括号用英文括号并加空格"),
    ("1.4", "符号：破折号用 em dash（—）"),
    ("1.5", "符号：中文间隔号 · 根据语义转换"),
    ("1.6", "符号：年份区间用 en dash（–），前后不空格"),
    ("1.7", "符号：括号嵌套时外圆内方"),
    ("2.1", "大小写：标题中所有实词首字母大写"),
    ("2.4", "大小写：头衔大小写规则"),
    ("2.5", "大小写：the X dynasty（dynasty 小写）"),
    ("2.6", "大小写：双引号内大小写三分支判断"),
    ("2.7", "大小写：Party（中共）一律大写"),
    ("2.8", "大小写：标题中介词一律小写"),
    ("2.9", "大小写：冒号后首字母大写"),
    ("3.1", "数字：0-9 用单词，10+ 用阿拉伯数字"),
    ("3.3", "数字：percentage point(s) 单复数"),
    ("3.4", "数字：日期格式 Month DD, YYYY"),
    ("3.5", "数字：金额加千分位，小数不超两位"),
    ("3.6", "数字：度量衡用全称"),
    ("3.9", "数字：句首不得用阿拉伯数字"),
    ("3.10", "数字：X odd 加 hyphen"),
    ("4.1", "专名：人名姓+名；吕=Lv"),
    ("4.5", "专名：机构名用官方英文译名"),
    ("4.9", "专名：永远用全称，不用缩写形式"),
    ("4.10.2", "专名：南粤=Guangdong；大陆=Chinese mainland"),
    ("5.1", "句法：无主语标题转被动；正文补主语"),
    ("5.2", "句法：截至…末 → By the end of…"),
    ("5.4", "句法：居第一 → ranked first in…"),
    ("6", "合并：并列同类专名合并，中心词小写"),
    ("7.1", "漏译：并列动词/名词均应译出"),
    ("7.2", "漏译：历史纪年与公元年均应译出"),
    ("7.3", "漏译：数字前限定词不应省略"),
    ("8.1–8.9", "理解：地理范围/程度副词/并列合并/修饰对象/同字异义等"),
    ("9.1–9.3", "句法优化：断句/年份位置/意译确切度"),
]
