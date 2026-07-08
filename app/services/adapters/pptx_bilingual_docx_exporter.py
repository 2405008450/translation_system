from __future__ import annotations

from collections import OrderedDict
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import re
from typing import Any, Iterable, Mapping

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from app.services.adapters.models import BlockNode, DocumentAST
from app.services.adapters.pptx_adapter import PptxAdapter
from app.services.document_exporter import DOCX_MEDIA_TYPE
from app.services.document_workspace import normalize_document_parse_options


PPTX_BILINGUAL_DOCX_EXPORT_TYPE = "bilingual_pptx_docx_layout"


@dataclass
class _BilingualCell:
    source_parts: list[str] = field(default_factory=list)
    target_parts: list[str] = field(default_factory=list)


@dataclass
class _BilingualItem:
    kind: str
    order: int
    section_key: str
    section_label: str
    title: str
    source_parts: list[str] = field(default_factory=list)
    target_parts: list[str] = field(default_factory=list)
    cells: dict[tuple[int, int], _BilingualCell] = field(default_factory=dict)


class PptxBilingualDocxExporter:
    """把 PPTX 按幻灯片阅读顺序导出为双语 Word 审校稿。"""

    def export(
        self,
        original_bytes: bytes,
        segments: Iterable[Any],
        filename: str,
        document_parse_options: Mapping[str, object] | str | None = None,
    ) -> tuple[bytes, str, str]:
        parse_options = normalize_document_parse_options(document_parse_options)
        parsed = PptxAdapter().parse_with_options(
            original_bytes,
            filename=filename or "source.pptx",
            options=parse_options,
        )
        translated_segments = _match_translated_segments(parsed.segments, segments)
        items = self._build_items(parsed.ast, parsed.segments, translated_segments)
        document = self._build_document(filename, items)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue(), DOCX_MEDIA_TYPE, build_bilingual_pptx_docx_filename(filename)

    def _build_items(
        self,
        ast: DocumentAST,
        parsed_segments: Iterable[Any],
        translated_segments: Mapping[str, Any],
    ) -> list[_BilingualItem]:
        items: "OrderedDict[tuple[Any, ...], _BilingualItem]" = OrderedDict()

        for order, parsed_segment in enumerate(parsed_segments):
            node = _resolve_node_by_path(ast, parsed_segment.block_path)
            if node is None:
                continue
            metadata = dict(node.metadata or {})
            key = _item_key(metadata, fallback_order=order)
            section_key, section_label = _section_context(metadata)
            item = items.get(key)
            if item is None:
                item = _BilingualItem(
                    kind=_item_kind(metadata),
                    order=order,
                    section_key=section_key,
                    section_label=section_label,
                    title=_item_title(metadata),
                )
                items[key] = item

            source_text = str(getattr(parsed_segment, "display_text", None) or parsed_segment.source_text or "")
            translated_segment = translated_segments.get(parsed_segment.segment_id)
            target_text = str(_get_segment_value(translated_segment, "target_text", "") or "")

            if item.kind == "table":
                cell_key = (
                    _to_int(metadata.get("row"), default=0),
                    _to_int(metadata.get("col"), default=0),
                )
                cell = item.cells.setdefault(cell_key, _BilingualCell())
                _append_text_part(cell.source_parts, source_text)
                _append_text_part(cell.target_parts, target_text)
            else:
                _append_text_part(item.source_parts, source_text)
                _append_text_part(item.target_parts, target_text)

        return list(items.values())

    def _build_document(self, filename: str, items: list[_BilingualItem]) -> Document:
        document = Document()
        _configure_document_page(document)

        title = document.add_heading("PPTX 双语审校稿", level=0)
        title.paragraph_format.space_after = Pt(4)
        subtitle = document.add_paragraph(Path(filename or "presentation.pptx").name)
        subtitle.runs[0].font.size = Pt(9)
        subtitle.runs[0].font.color.rgb = RGBColor(90, 90, 90)

        sections: "OrderedDict[str, tuple[str, list[_BilingualItem]]]" = OrderedDict()
        for item in items:
            if item.section_key not in sections:
                sections[item.section_key] = (item.section_label, [])
            sections[item.section_key][1].append(item)

        for section_index, (_, (section_label, section_items)) in enumerate(sections.items()):
            if section_index > 0:
                document.add_page_break()
            document.add_heading(section_label, level=1)
            self._append_section_items(document, section_items)

        return document

    def _append_section_items(self, document: Document, items: list[_BilingualItem]) -> None:
        text_table = None

        for item in sorted(items, key=lambda value: value.order):
            if item.kind == "table":
                text_table = None
                _append_original_table(document, item)
                continue

            if text_table is None:
                text_table = _create_bilingual_text_table(document)
            _append_bilingual_text_row(
                text_table,
                source_text=_join_text_parts(item.source_parts),
                target_text=_join_text_parts(item.target_parts),
                label=item.title,
            )


def build_bilingual_pptx_docx_filename(filename: str) -> str:
    source_path = Path(filename or "presentation.pptx")
    stem = source_path.stem or "presentation"
    return f"{stem}_pptx_bilingual.docx"


def _configure_document_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)


def _create_bilingual_text_table(document: Document):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _write_header_cell(table.rows[0].cells[0], "原文")
    _write_header_cell(table.rows[0].cells[1], "译文")
    return table


def _append_bilingual_text_row(table, *, source_text: str, target_text: str, label: str) -> None:
    row = table.add_row()
    _write_bilingual_cell(row.cells[0], source_text, label=label)
    _write_bilingual_cell(row.cells[1], target_text)


def _append_original_table(document: Document, item: _BilingualItem) -> None:
    document.add_paragraph(item.title)
    max_row = max((row for row, _ in item.cells), default=-1)
    max_col = max((col for _, col in item.cells), default=-1)
    if max_row < 0 or max_col < 0:
        return

    table = document.add_table(rows=max_row + 1, cols=max_col + 1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for (row_index, col_index), cell_item in item.cells.items():
        cell = table.rows[row_index].cells[col_index]
        _write_bilingual_cell(
            cell,
            _join_text_parts(cell_item.source_parts),
            target_text=_join_text_parts(cell_item.target_parts),
        )


def _write_header_cell(cell, text: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = True


def _write_bilingual_cell(
    cell,
    source_text: str,
    *,
    target_text: str | None = None,
    label: str = "",
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    paragraph = cell.paragraphs[0]

    if label:
        label_run = paragraph.add_run(f"{label}\n")
        label_run.bold = True
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(90, 90, 90)

    source_run = paragraph.add_run(source_text or "")
    source_run.font.size = Pt(10)

    if target_text is None:
        return

    target_paragraph = cell.add_paragraph()
    target_run = target_paragraph.add_run(target_text or "（未翻译）")
    target_run.font.size = Pt(10)
    target_run.font.color.rgb = RGBColor(31, 95, 175)
    if not target_text:
        target_run.italic = True


def _append_text_part(parts: list[str], text: str) -> None:
    normalized = str(text or "").strip()
    if normalized:
        parts.append(normalized)


def _join_text_parts(parts: Iterable[str]) -> str:
    values = [part.strip() for part in parts if str(part or "").strip()]
    if not values:
        return ""
    separator = "\n" if any("\n" in value for value in values) else " "
    return separator.join(values)


def _item_key(metadata: Mapping[str, Any], *, fallback_order: int) -> tuple[Any, ...]:
    item_type = str(metadata.get("item_type") or "")
    part_name = str(metadata.get("part_name") or "")
    if item_type == "table_cell":
        return ("table", part_name, _to_int(metadata.get("table_index"), default=0))
    if item_type == "paragraph":
        return ("paragraph", part_name, _to_int(metadata.get("paragraph_index"), default=fallback_order))
    if item_type == "chart_text":
        return (
            "chart_text",
            part_name,
            str(metadata.get("chart_text_kind") or ""),
            _to_int(metadata.get("chart_text_index"), default=fallback_order),
        )
    if item_type == "comment":
        return ("comment", part_name, _to_int(metadata.get("comment_index"), default=fallback_order))
    if item_type == "document_property":
        return ("document_property", part_name, _to_int(metadata.get("property_index"), default=fallback_order))
    return ("item", part_name, fallback_order)


def _item_kind(metadata: Mapping[str, Any]) -> str:
    if metadata.get("item_type") == "table_cell":
        return "table"
    return "text"


def _item_title(metadata: Mapping[str, Any]) -> str:
    item_type = str(metadata.get("item_type") or "")
    part_kind = str(metadata.get("part_kind") or "")
    if item_type == "table_cell":
        return f"表格 {_to_int(metadata.get('table_index'), default=0) + 1}"
    if part_kind == "notes":
        return "演讲者备注"
    if item_type == "chart_text":
        return "图表文本"
    if item_type == "comment":
        return "批注"
    if item_type == "document_property":
        property_name = str(metadata.get("property_name") or "").strip()
        return f"文档属性：{property_name}" if property_name else "文档属性"
    return "文本框"


def _section_context(metadata: Mapping[str, Any]) -> tuple[str, str]:
    item_type = str(metadata.get("item_type") or "")
    part_kind = str(metadata.get("part_kind") or "")
    part_name = str(metadata.get("part_name") or "")

    if part_kind == "slide":
        slide_number = _part_number(part_name, "slide")
        return part_name, f"幻灯片 {slide_number}" if slide_number else "幻灯片"
    if part_kind == "notes":
        notes_number = _part_number(part_name, "notesSlide")
        return part_name, f"演讲者备注 {notes_number}" if notes_number else "演讲者备注"
    if part_kind == "chart" or item_type == "chart_text":
        chart_number = _part_number(part_name, "chart")
        return part_name, f"图表 {chart_number}" if chart_number else "图表"
    if item_type == "comment":
        return "ppt-comments", "批注"
    if item_type == "document_property":
        return "ppt-document-properties", "文档属性"
    return part_name or "pptx", "PPTX 文本"


def _part_number(part_name: str, prefix: str) -> str:
    match = re.search(rf"{re.escape(prefix)}(\d+)\.xml$", part_name)
    return match.group(1) if match else ""


def _match_translated_segments(parsed_segments: Iterable[Any], translated_segments: Iterable[Any]) -> dict[str, Any]:
    parsed_list = list(parsed_segments)
    translated_list = list(translated_segments)
    translated_by_id = {
        str(_get_segment_value(segment, "sentence_id", _get_segment_value(segment, "segment_id", ""))): segment
        for segment in translated_list
    }
    translated_by_source: dict[str, list[Any]] = {}
    for segment in translated_list:
        source_key = _normalize_source_text(_get_segment_value(segment, "source_text", ""))
        if source_key:
            translated_by_source.setdefault(source_key, []).append(segment)

    matched: dict[str, Any] = {}
    used: set[int] = set()
    for index, parsed_segment in enumerate(parsed_list):
        parsed_id = str(parsed_segment.segment_id)
        parsed_source = _normalize_source_text(parsed_segment.source_text)
        candidate = translated_by_id.get(parsed_id)
        if candidate is not None and _normalize_source_text(_get_segment_value(candidate, "source_text", "")) != parsed_source:
            candidate = None
        if candidate is None and index < len(translated_list):
            indexed = translated_list[index]
            if id(indexed) not in used and _normalize_source_text(_get_segment_value(indexed, "source_text", "")) == parsed_source:
                candidate = indexed
        if candidate is None and parsed_source:
            candidates = translated_by_source.get(parsed_source, [])
            while candidates and id(candidates[0]) in used:
                candidates.pop(0)
            if candidates:
                candidate = candidates.pop(0)
        if candidate is not None:
            matched[parsed_id] = candidate
            used.add(id(candidate))
    return matched


def _resolve_node_by_path(ast: DocumentAST, block_path: str) -> BlockNode | None:
    if not block_path:
        return None
    parts = block_path.split(".")
    try:
        node = ast.nodes[int(parts[0])]
    except (IndexError, TypeError, ValueError):
        return None
    index = 1
    while index < len(parts):
        if parts[index] != "children" or index + 1 >= len(parts):
            return None
        try:
            node = node.children[int(parts[index + 1])]
        except (IndexError, TypeError, ValueError):
            return None
        index += 2
    return node


def _get_segment_value(segment: Any, field_name: str, default: Any = None) -> Any:
    if segment is None:
        return default
    if isinstance(segment, Mapping):
        return segment.get(field_name, default)
    return getattr(segment, field_name, default)


def _normalize_source_text(value: Any) -> str:
    return " ".join(str(value or "").split())


def _to_int(value: Any, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default
