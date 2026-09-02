from __future__ import annotations

import json
import logging
import mimetypes
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import Any
import re

import yaml

from app.services.adapters.export_formats import get_supported_exports
from app.services.adapters.tmx_exporter import TmxExporter
from app.services.adapters.xliff_exporter import XliffExporter


logger = logging.getLogger(__name__)


class MultiFormatExporter:
    def __init__(
        self,
        source_lang: str = "zh-CN",
        target_lang: str = "en-US",
    ):
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.tmx_exporter = TmxExporter(source_lang, target_lang)
        self.xliff_exporter = XliffExporter(source_lang, target_lang)

    def get_available_exports(self, filename: str) -> list[dict]:
        extension = Path(filename).suffix.lower()
        return [
            {
                "id": option.id,
                "name": option.name,
                "description": option.description,
                "extension": option.extension or extension,
            }
            for option in get_supported_exports(extension)
        ]

    def export(
        self,
        export_type: str,
        segments: list[Any],
        filename: str,
        original_bytes: bytes | None = None,
    ) -> tuple[bytes, str, str]:
        normalized_segments = self._normalize_segments(segments)
        translation_maps = self._build_translation_maps(normalized_segments)
        extension = Path(filename).suffix.lower()
        base_name = Path(filename).stem or "translated"

        if export_type == "source":
            if original_bytes is None:
                raise ValueError("Source export requires the original source file.")
            return original_bytes, self._get_mime_type(extension), filename
        if export_type == "original":
            return self._export_original(
                extension,
                filename,
                original_bytes,
                normalized_segments,
                translation_maps,
            )
        if export_type == "bilingual":
            return self._export_bilingual_original(
                extension,
                base_name,
                normalized_segments,
                original_bytes,
                translation_maps,
            )
        if export_type == "bilingual_docx":
            return self._export_bilingual_docx(normalized_segments, base_name)
        if export_type == "bilingual_txt":
            return self._export_bilingual_txt(normalized_segments, base_name)
        if export_type == "bilingual_excel_original":
            return self._export_bilingual_excel_original(
                extension,
                normalized_segments,
                base_name,
                original_bytes,
            )
        if export_type == "bilingual_excel":
            return self._export_bilingual_excel(normalized_segments, base_name)
        if export_type == "tmx":
            return self._export_tmx(normalized_segments, base_name)
        if export_type in {"xliff", "xliff2"}:
            version = "2.0" if export_type == "xliff2" else "1.2"
            return self._export_xliff(normalized_segments, filename, version)

        raise ValueError(f"Unsupported export type: {export_type}")

    def _normalize_segments(self, segments: list[Any]) -> list[dict[str, Any]]:
        normalized: list[dict[str, Any]] = []
        for index, segment in enumerate(segments):
            if isinstance(segment, dict):
                item = dict(segment)
            else:
                item = {
                    "segment_id": getattr(segment, "segment_id", None),
                    "sentence_id": getattr(segment, "sentence_id", None),
                    "source_text": getattr(segment, "source_text", ""),
                    "display_text": getattr(segment, "display_text", ""),
                    "target_text": getattr(segment, "target_text", ""),
                    "status": getattr(segment, "status", "none"),
                    "matched_source_text": getattr(segment, "matched_source_text", ""),
                    "block_type": getattr(segment, "block_type", "paragraph"),
                    "block_index": getattr(segment, "block_index", index),
                    "row_index": getattr(segment, "row_index", None),
                    "cell_index": getattr(segment, "cell_index", None),
                }

            sentence_id = item.get("sentence_id") or item.get("segment_id") or f"seg_{index}"
            item.setdefault("segment_id", sentence_id)
            item.setdefault("sentence_id", sentence_id)
            item.setdefault("source_text", "")
            item.setdefault("display_text", item.get("source_text", ""))
            item.setdefault("target_text", "")
            item.setdefault("status", "none")
            item.setdefault("matched_source_text", "")
            item.setdefault("block_type", "paragraph")
            item.setdefault("block_index", index)
            item.setdefault("row_index", None)
            item.setdefault("cell_index", None)
            normalized.append(item)
        return normalized

    def _build_translation_maps(self, segments: list[dict[str, Any]]) -> dict[str, dict[str, str]]:
        source_candidates: dict[str, set[str]] = defaultdict(set)
        display_candidates: dict[str, set[str]] = defaultdict(set)
        key_map: dict[str, str] = {}
        path_map: dict[str, str] = {}
        row_col_map: dict[str, str] = {}
        index_map: dict[str, str] = {}
        segment_id_map: dict[str, str] = {}

        for segment in segments:
            target_text = str(segment.get("target_text") or "").strip()
            if not target_text:
                continue

            source_text = str(segment.get("source_text") or "").strip()
            display_text = str(segment.get("display_text") or "").strip()
            segment_id = str(segment.get("segment_id") or segment.get("sentence_id") or "")
            if segment_id:
                segment_id_map[segment_id] = target_text
            if source_text:
                source_candidates[source_text].add(target_text)
            if display_text:
                display_candidates[display_text].add(target_text)

            key = str(segment.get("key") or "").strip()
            if key:
                key_map[key] = target_text

            metadata_path = str(
                segment.get("metadata_path")
                or segment.get("json_path")
                or segment.get("path")
                or ""
            ).strip()
            if metadata_path:
                path_map[metadata_path] = target_text

            row_index = self._to_optional_int(segment.get("row_index"))
            cell_index = self._to_optional_int(segment.get("cell_index"))
            if row_index is not None and cell_index is not None:
                row_col_map[f"{row_index},{cell_index}"] = target_text

            subtitle_index = segment.get("subtitle_index", segment.get("index"))
            if subtitle_index is not None and str(subtitle_index).strip():
                index_map[str(subtitle_index).strip()] = target_text

        return {
            "source_text": self._collapse_unique_values(source_candidates),
            "display_text": self._collapse_unique_values(display_candidates),
            "key": key_map,
            "path": path_map,
            "row_col": row_col_map,
            "index": index_map,
            "segment_id": segment_id_map,
        }

    def _collapse_unique_values(self, candidates: dict[str, set[str]]) -> dict[str, str]:
        result: dict[str, str] = {}
        for key, values in candidates.items():
            if len(values) == 1:
                result[key] = next(iter(values))
            elif len(values) > 1:
                # 有多个不同译文时，使用第一个（按字母顺序）
                # 这比完全跳过更好，至少能保证有翻译输出
                result[key] = sorted(values)[0]
        return result

    def _export_original(
        self,
        extension: str,
        filename: str,
        original_bytes: bytes | None,
        segments: list[dict[str, Any]],
        translation_maps: dict[str, dict[str, str]],
    ) -> tuple[bytes, str, str]:
        if original_bytes is None:
            raise ValueError("Original export requires the original source file.")

        if extension in {".dwg", ".dxf"}:
            segments = self._collapse_cad_sentence_segments(segments)
            translation_maps = self._build_translation_maps(segments)

        text_map = self._build_text_translation_map(translation_maps)
        if extension in {".dwg", ".dxf"}:
            # 普通 source->target 映射无法可靠表达同一 MTEXT 内的多段译文；
            # 额外按 handle/段落序号重建整块译文，让导出器一次性整体回写。
            text_map.update(self._build_cad_mtext_translation_map(segments))
        export_filename = self._build_translated_filename(filename)

        if extension in {".txt", ".dat"}:
            content = self._replace_plain_text(original_bytes, text_map)
        elif extension in {".html", ".htm"}:
            from app.services.adapters.html_exporter import HtmlExporter

            content = HtmlExporter().export_by_segments(original_bytes, segments, text_map)
        elif extension in {".md", ".markdown"}:
            from app.services.adapters.markdown_exporter import MarkdownExporter

            content = MarkdownExporter().export(original_bytes, text_map)
        elif extension == ".csv":
            from app.services.adapters.csv_exporter import CsvExporter

            content = CsvExporter().export(
                original_bytes,
                {**text_map, **translation_maps["row_col"]},
            )
        elif extension == ".properties":
            from app.services.adapters.properties_exporter import PropertiesExporter

            content = PropertiesExporter().export(
                original_bytes,
                {**text_map, **translation_maps["key"]},
            )
        elif extension in {".po", ".pot"}:
            from app.services.adapters.po_exporter import PoExporter

            content = PoExporter().export(original_bytes, text_map)
        elif extension == ".strings":
            from app.services.adapters.strings_exporter import StringsExporter

            content = StringsExporter().export(
                original_bytes,
                {**text_map, **translation_maps["key"]},
            )
        elif extension == ".srt":
            from app.services.adapters.srt_exporter import SrtExporter

            content = SrtExporter().export(
                original_bytes,
                {**text_map, **translation_maps["index"]},
            )
        elif extension == ".json":
            content = self._export_json(original_bytes, text_map, translation_maps["path"])
        elif extension in {".yaml", ".yml"}:
            content = self._export_yaml(original_bytes, text_map, translation_maps["path"])
        elif extension == ".php":
            content = self._replace_plain_text(original_bytes, text_map)
        elif extension in {".dita", ".ditamap", ".xml"}:
            from app.services.adapters.dita_exporter import DitaExporter

            content = DitaExporter().export_with_translations(original_bytes, text_map)
        elif extension == ".svg":
            from app.services.adapters.svg_exporter import SvgExporter

            content = SvgExporter().export(original_bytes, text_map)
        elif extension == ".pptx":
            from app.services.adapters.pptx_exporter import PptxExporter

            content = PptxExporter().export(original_bytes, segments)
        elif extension == ".xlsx":
            from app.services.adapters.xlsx_exporter import XlsxExporter

            content = XlsxExporter().export(original_bytes, segments)
        elif extension == ".sdlxliff":
            from app.services.adapters.sdlxliff_exporter import SdlxliffExporter

            content = SdlxliffExporter().export_by_segments(original_bytes, segments)
        elif extension == ".txml":
            from app.services.adapters.txml_exporter import TxmlExporter

            content = TxmlExporter().export(original_bytes, text_map)
        elif extension == ".dxf":
            from app.services.adapters.dxf_exporter import DxfExporter, DxfExportOptions

            # 从 segments 中提取合并文本信息
            merged_text_info = self._extract_merged_text_info(segments, text_map)
            mtext_split_info = self._extract_mtext_split_info(segments, text_map)
            has_merged_groups = bool(merged_text_info)
            
            dxf_options = DxfExportOptions(
                enable_overflow_shrink=True,
                enable_spatial_merge_export=has_merged_groups,
            )
            content = DxfExporter().export(
                original_bytes, 
                text_map,
                options=dxf_options,
                merged_text_info=merged_text_info,
                mtext_split_info=mtext_split_info,
            )
        elif extension == ".dwg":
            from app.services.adapters.dwg_exporter import DwgExporter

            # 从 segments 中提取合并文本信息
            merged_text_info = self._extract_merged_text_info(segments, text_map)
            mtext_split_info = self._extract_mtext_split_info(segments, text_map)

            result = DwgExporter().export_with_extension(
                original_bytes, 
                text_map,
                merged_text_info=merged_text_info,
                mtext_split_info=mtext_split_info,
            )
            content = result.content
            # 当 DWG 回写不可用时降级为 DXF，需要纠正下游 mime/扩展
            if result.extension != ".dwg":
                extension = result.extension
                export_filename = self._build_translated_filename(
                    filename, extension_override=result.extension
                )
        elif extension == ".idml":
            from app.services.adapters.idml_exporter import IdmlExporter

            content = IdmlExporter().export(original_bytes, text_map)
        elif extension == ".mif":
            from app.services.adapters.mif_exporter import MifExporter

            content = MifExporter().export(original_bytes, text_map)
        elif extension == ".zip":
            from app.services.adapters.zip_exporter import ZipExporter

            content = ZipExporter().export(original_bytes, text_map, segments=segments)
        elif extension == ".rar":
            from app.services.adapters.rar_exporter import RarExporter

            content = RarExporter().export(original_bytes, text_map, segments=segments)
        else:
            raise ValueError(f"Original export is not supported for {extension}.")

        return content, self._get_mime_type(extension), export_filename

    def _export_bilingual_original(
        self,
        extension: str,
        base_name: str,
        segments: list[dict[str, Any]],
        original_bytes: bytes | None,
        translation_maps: dict[str, dict[str, str]],
    ) -> tuple[bytes, str, str]:
        if original_bytes is None:
            raise ValueError("Bilingual export requires the original source file.")

        if extension in {".properties", ".po", ".pot", ".strings", ".html", ".htm", ".srt"}:
            if extension == ".properties":
                content = self._export_bilingual_properties(original_bytes, segments)
            elif extension in {".po", ".pot"}:
                content = self._export_bilingual_po(original_bytes, translation_maps["source_text"])
            elif extension == ".strings":
                content = self._export_bilingual_strings(original_bytes, segments)
            elif extension in {".html", ".htm"}:
                content = self._export_bilingual_html(original_bytes, segments)
            else:
                content = self._export_bilingual_srt(original_bytes, segments)
            return content, self._get_mime_type(extension), f"{base_name}-bilingual{extension}"

        return self._export_bilingual_txt(segments, base_name)

    def _export_bilingual_properties(
        self,
        original_bytes: bytes,
        segments: list[dict[str, Any]],
    ) -> bytes:
        from app.services.adapters.properties_exporter import PropertiesExporter

        exporter = PropertiesExporter()
        content = exporter._decode_content(original_bytes)
        lines = content.replace("\r\n", "\n").split("\n")
        source_to_target = self._build_source_to_target_map(segments)
        result_lines: list[str] = []
        for line in lines:
            if not line.strip() or line.lstrip().startswith(("#", "!")):
                result_lines.append(line)
                continue
            key, value, separator = exporter._parse_line(line)
            clean_value = value.strip()
            target = source_to_target.get(clean_value, "")
            bilingual_value = clean_value if not target else f"{clean_value} | {target}"
            result_lines.append(f"{key}{separator}{exporter._escape_value(bilingual_value)}")
        return "\n".join(result_lines).encode("utf-8")

    def _export_bilingual_po(
        self,
        original_bytes: bytes,
        translations: dict[str, str],
    ) -> bytes:
        from app.services.adapters.po_exporter import PoExporter

        return PoExporter().export(original_bytes, translations)

    def _export_bilingual_strings(
        self,
        original_bytes: bytes,
        segments: list[dict[str, Any]],
    ) -> bytes:
        from app.services.adapters.strings_exporter import StringsExporter

        exporter = StringsExporter()
        content = exporter._decode_content(original_bytes)
        source_to_target = self._build_source_to_target_map(segments)

        def replace_value(match):
            original_value = exporter._unescape(match.group(2))
            target = source_to_target.get(original_value, "")
            bilingual = original_value if not target else f"{original_value} | {target}"
            return f'"{match.group(1)}" = "{exporter._escape(bilingual)}";'

        pattern = r'"([^"\\]*(?:\\.[^"\\]*)*)"\s*=\s*"([^"\\]*(?:\\.[^"\\]*)*)"\s*;'
        return re.sub(pattern, replace_value, content).encode("utf-8")

    def _export_bilingual_html(
        self,
        original_bytes: bytes,
        segments: list[dict[str, Any]],
    ) -> bytes:
        import re

        from app.services.adapters.html_exporter import HtmlExporter

        exporter = HtmlExporter()
        content = exporter._decode_content(original_bytes)
        source_to_target = self._build_source_to_target_map(segments)

        # 标记 <title> 区域，避免在其中插入双语标记
        _SKIP_BILINGUAL_RE = re.compile(
            r'(<title[^>]*>)(.*?)(</title>)', re.IGNORECASE | re.DOTALL
        )

        def replace_text_node(match):
            text = match.group(1)
            stripped = text.strip()
            target = source_to_target.get(stripped, "")
            if not target:
                return match.group(0)
            leading = text[: len(text) - len(text.lstrip())]
            trailing = text[len(text.rstrip()) :]
            # 原文保持原样，译文紧跟其后，用不同颜色区分
            bilingual = (
                f'{leading}'
                f'<span style="color:#333;">{stripped}</span>'
                f'<br/>'
                f'<span style="color:#0066cc; font-style:italic;">{target}</span>'
                f'{trailing}'
            )
            return f">{bilingual}<"

        # 先保护 <title> 内容不被替换
        title_placeholders: list[str] = []

        def protect_title(m):
            placeholder = f"__TITLE_PLACEHOLDER_{len(title_placeholders)}__"
            title_placeholders.append(m.group(0))
            return placeholder

        protected = _SKIP_BILINGUAL_RE.sub(protect_title, content)

        result = re.sub(r">([^<]+)<", replace_text_node, protected)

        # 恢复 <title> 内容
        for i, original_title in enumerate(title_placeholders):
            result = result.replace(f"__TITLE_PLACEHOLDER_{i}__", original_title)

        # 注入双语对照的样式
        bilingual_style = (
            '<style>\n'
            '.bilingual-source { color: #333; display: block; margin-bottom: 2px; }\n'
            '.bilingual-target { color: #0066cc; font-style: italic; display: block; '
            'margin-bottom: 8px; padding-left: 0; }\n'
            '</style>'
        )
        head_close = re.search(r'</head>', result, re.IGNORECASE)
        if head_close:
            result = result[:head_close.start()] + '\n' + bilingual_style + '\n' + result[head_close.start():]
        else:
            body_open = re.search(r'<body[^>]*>', result, re.IGNORECASE)
            if body_open:
                result = result[:body_open.start()] + bilingual_style + '\n' + result[body_open.start():]
            else:
                result = bilingual_style + '\n' + result

        result = exporter._normalize_fonts(result)
        result = exporter._ensure_utf8_charset(result)
        return result.encode("utf-8")

    def _export_bilingual_srt(
        self,
        original_bytes: bytes,
        segments: list[dict[str, Any]],
    ) -> bytes:
        import re

        from app.services.adapters.srt_exporter import TIMECODE_PATTERN, SrtExporter

        exporter = SrtExporter()
        content = exporter._decode_content(original_bytes).replace("\r\n", "\n").replace("\r", "\n")
        source_to_target = self._build_source_to_target_map(segments)
        blocks = re.split(r"\n\n+", content.strip())
        result_blocks: list[str] = []
        for block in blocks:
            lines = block.strip().split("\n")
            if len(lines) < 2 or not TIMECODE_PATTERN.match(lines[1]):
                result_blocks.append(block)
                continue
            original_text = "\n".join(lines[2:])
            clean_text = re.sub(r"<[^>]+>", "", original_text).strip()
            target = source_to_target.get(clean_text, "")
            if target:
                result_blocks.append(f"{lines[0]}\n{lines[1]}\n{original_text}\n{target}")
            else:
                result_blocks.append(block)
        return "\n\n".join(result_blocks).encode("utf-8")

    def _export_bilingual_docx(
        self,
        segments: list[dict[str, Any]],
        base_name: str,
    ) -> tuple[bytes, str, str]:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT

        document = Document()
        document.add_heading("双语对照文档", level=1)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "原文"
        table.rows[0].cells[1].text = "译文"

        for segment in segments:
            source_text = str(segment.get("source_text") or "")
            if not source_text:
                continue
            target_text = str(segment.get("target_text") or "")
            row = table.add_row()
            row.cells[0].text = source_text
            row.cells[1].text = target_text

        buffer = BytesIO()
        document.save(buffer)
        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{base_name}-bilingual.docx",
        )

    def _export_bilingual_txt(
        self,
        segments: list[dict[str, Any]],
        base_name: str,
    ) -> tuple[bytes, str, str]:
        lines: list[str] = []
        for index, segment in enumerate(segments, start=1):
            source_text = str(segment.get("source_text") or "")
            if not source_text:
                continue
            target_text = str(segment.get("target_text") or "")
            lines.append(f"[{index}] 原文: {source_text}")
            lines.append(f"[{index}] 译文: {target_text or '(未翻译)'}")
            lines.append("")

        return (
            "\n".join(lines).encode("utf-8"),
            "text/plain; charset=utf-8",
            f"{base_name}-bilingual.txt",
        )

    def _export_bilingual_excel(
        self,
        segments: list[dict[str, Any]],
        base_name: str,
    ) -> tuple[bytes, str, str]:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        wb = Workbook()
        ws = wb.active
        ws.title = "双语对照"

        # 表头样式
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(fill_type="solid", fgColor="4472C4")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        ws.column_dimensions["A"].width = 50
        ws.column_dimensions["B"].width = 50

        ws.append(["原文", "译文"])
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        cell_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        for segment in segments:
            source_text = str(segment.get("source_text") or "")
            if not source_text:
                continue
            target_text = str(segment.get("target_text") or "")
            row_idx = ws.max_row + 1
            ws.append([source_text, target_text])
            for cell in ws[row_idx]:
                cell.alignment = cell_alignment

        buffer = BytesIO()
        wb.save(buffer)
        return (
            buffer.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{base_name}-bilingual.xlsx",
        )

    def _export_bilingual_excel_original(
        self,
        extension: str,
        segments: list[dict[str, Any]],
        base_name: str,
        original_bytes: bytes | None,
    ) -> tuple[bytes, str, str]:
        if extension != ".xlsx":
            raise ValueError("Original-format bilingual Excel export requires an XLSX source file.")
        if original_bytes is None:
            raise ValueError("Original-format bilingual Excel export requires the original source file.")

        from app.services.adapters.xlsx_exporter import XlsxExporter

        return (
            XlsxExporter().export(original_bytes, segments, bilingual=True),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            f"{base_name}_bilingual.xlsx",
        )

    def _export_tmx(
        self,
        segments: list[dict[str, Any]],
        base_name: str,
    ) -> tuple[bytes, str, str]:
        content = self.tmx_exporter.export(segments, base_name)
        return content, "application/x-tmx+xml", f"{base_name}.tmx"

    def _export_xliff(
        self,
        segments: list[dict[str, Any]],
        filename: str,
        version: str,
    ) -> tuple[bytes, str, str]:
        extension = Path(filename).suffix.lower()
        format_map = {
            ".docx": "winword",
            ".pdf": "pdf",
            ".pptx": "powerpoint",
            ".xlsx": "x-excel",
            ".txt": "plaintext",
            ".dat": "plaintext",
            ".html": "html",
            ".htm": "html",
            ".xml": "xml",
            ".dita": "xml",
            ".json": "json",
            ".yaml": "yaml",
            ".yml": "yaml",
        }
        self.xliff_exporter.version = version
        content = self.xliff_exporter.export(
            segments,
            filename,
            format_map.get(extension, "plaintext"),
        )
        return content, "application/xliff+xml", f"{Path(filename).stem or 'translated'}.xlf"

    def _build_text_translation_map(
        self,
        translation_maps: dict[str, dict[str, str]],
    ) -> dict[str, str]:
        return {
            **translation_maps["display_text"],
            **translation_maps["source_text"],
        }

    def _build_cad_mtext_translation_map(
        self,
        segments: list[dict[str, Any]],
    ) -> dict[str, str]:
        """按 MTEXT 稳定身份重建整块译文，避免逐段匹配造成中英混排。"""
        from app.services.adapters.dxf_adapter import clean_mtext
        from app.services.adapters.dxf_exporter import (
            CAD_MTEXT_HANDLE_BLOCK_PREFIX,
            CAD_MTEXT_HANDLE_TRANSLATION_PREFIX,
            CAD_MTEXT_SOURCE_BLOCK_PREFIX,
        )

        groups: dict[tuple[str, str, str], list[dict[str, Any]]] = defaultdict(list)
        incomplete_groups: set[tuple[str, str, str]] = set()

        for segment in segments:
            metadata = segment.get("metadata", {}) or {}
            if str(metadata.get("entity_type") or "").upper() != "MTEXT":
                continue
            if "mtext_para_index" not in metadata:
                continue
            handle = str(metadata.get("handle") or "")
            raw = str(metadata.get("mtext_raw") or "")
            if not handle or not raw:
                continue
            key = (str(metadata.get("scope") or ""), handle, raw)
            groups[key].append(segment)
            if not str(segment.get("target_text") or "").strip():
                incomplete_groups.add(key)

        candidates: dict[str, set[str]] = defaultdict(set)
        handle_translations: dict[str, str] = {}
        blocked_sources: set[str] = set()
        for key, items in groups.items():
            if key in incomplete_groups:
                cleaned_source = clean_mtext(key[2])
                handle_translations[CAD_MTEXT_HANDLE_BLOCK_PREFIX + key[1]] = "1"
                handle_translations[CAD_MTEXT_SOURCE_BLOCK_PREFIX + cleaned_source] = "1"
                blocked_sources.add(cleaned_source)
                continue
            cleaned = clean_mtext(key[2])
            source_parts = cleaned.split("\n")
            targets_by_paragraph: dict[int, list[dict[str, Any]]] = defaultdict(list)
            valid_group = True

            for item in items:
                metadata = item.get("metadata", {}) or {}
                try:
                    paragraph_index = int(metadata.get("mtext_para_index"))
                except (TypeError, ValueError):
                    valid_group = False
                    break
                if paragraph_index < 0 or paragraph_index >= len(source_parts):
                    valid_group = False
                    break
                targets_by_paragraph[paragraph_index].append(item)

            if not valid_group or not targets_by_paragraph:
                continue

            # 只有所有非空原段都找到了对应句段，才允许整块重建；否则宁可交给
            # 后续安全兜底，也不能把某一段译文误写到另一段的位置。
            expected_paragraphs = {
                index for index, source_part in enumerate(source_parts) if source_part.strip()
            }
            if not expected_paragraphs.issubset(targets_by_paragraph.keys()):
                continue

            target_parts = list(source_parts)
            for paragraph_index, paragraph_items in targets_by_paragraph.items():
                paragraph_items.sort(key=lambda item: int(
                    (item.get("metadata", {}) or {}).get("cad_sentence_index", 0) or 0
                ))
                first_metadata = paragraph_items[0].get("metadata", {}) or {}
                joiner = str(first_metadata.get("cad_sentence_joiner", " "))
                paragraph_target = joiner.join(
                    str(item.get("target_text") or "").strip()
                    for item in paragraph_items
                ).strip()
                if not paragraph_target:
                    valid_group = False
                    break
                target_parts[paragraph_index] = paragraph_target

            if not valid_group:
                continue
            target_text = "\\P".join(target_parts)
            handle_translations[
                CAD_MTEXT_HANDLE_TRANSLATION_PREFIX + key[1]
            ] = target_text
            candidates[cleaned].add(target_text)

        # handle 映射用于精确回写；文本映射仅作为 DWG 重转导致 handle 变化时的
        # 唯一译文兜底。相同源文本存在多个译法时不得任选其一。
        handle_translations.update({
            source_text: next(iter(target_values))
            for source_text, target_values in candidates.items()
            if len(target_values) == 1 and source_text not in blocked_sources
        })
        return handle_translations

    def _build_source_to_target_map(self, segments: list[dict[str, Any]]) -> dict[str, str]:
        result: dict[str, str] = {}
        for segment in segments:
            source_text = str(segment.get("source_text") or "").strip()
            target_text = str(segment.get("target_text") or "").strip()
            if source_text and target_text and source_text not in result:
                result[source_text] = target_text
        return result

    def _replace_plain_text(self, original_bytes: bytes, translations: dict[str, str]) -> bytes:
        content = self._decode_text_content(original_bytes)

        # 构建规范化文本 -> 译文的映射（处理 source_text 被 normalize 过的情况）
        normalized_map: dict[str, str] = {}
        for source_text, target_text in translations.items():
            normalized_map[source_text] = target_text
            # 也用规范化后的 key 存一份，方便后续匹配
            normalized_key = re.sub(r'\s+', ' ', source_text.strip())
            if normalized_key != source_text:
                normalized_map[normalized_key] = target_text

        # 统一换行符
        unified = content.replace("\r\n", "\n").replace("\r", "\n")

        # 按空行分割段落（与 TxtAdapter._split_paragraphs 逻辑一致）
        parts = re.split(r'(\n\s*\n)', unified)

        result_parts: list[str] = []
        for part in parts:
            # 如果是段落分隔符（空行），保留原样
            if re.match(r'^\n\s*\n$', part):
                result_parts.append(part)
                continue

            stripped = part.strip()
            if not stripped:
                result_parts.append(part)
                continue

            # 将段落文本规范化后尝试匹配
            normalized_paragraph = re.sub(r'\s+', ' ', stripped)

            if normalized_paragraph in normalized_map:
                # 整段匹配成功，替换为译文
                result_parts.append(normalized_map[normalized_paragraph])
            else:
                # 尝试按句子级别替换（段落内可能有多个句子）
                replaced = part
                for source_text in sorted(normalized_map.keys(), key=len, reverse=True):
                    replaced = replaced.replace(source_text, normalized_map[source_text])
                result_parts.append(replaced)

        return "".join(result_parts).encode("utf-8")

    def _export_json(
        self,
        original_bytes: bytes,
        text_map: dict[str, str],
        path_map: dict[str, str],
    ) -> bytes:
        payload = json.loads(self._decode_text_content(original_bytes))
        translated = self._translate_tree(payload, text_map, path_map)
        return json.dumps(translated, ensure_ascii=False, indent=2).encode("utf-8")

    def _export_yaml(
        self,
        original_bytes: bytes,
        text_map: dict[str, str],
        path_map: dict[str, str],
    ) -> bytes:
        payload = yaml.safe_load(self._decode_text_content(original_bytes))
        translated = self._translate_tree(payload, text_map, path_map)
        return yaml.safe_dump(
            translated,
            allow_unicode=True,
            default_flow_style=False,
            sort_keys=False,
        ).encode("utf-8")

    def _translate_tree(
        self,
        payload: Any,
        text_map: dict[str, str],
        path_map: dict[str, str],
        path_parts: list[str] | None = None,
    ) -> Any:
        current_path_parts = path_parts or []
        current_path = "".join(current_path_parts) if current_path_parts and current_path_parts[0].startswith("[") else ".".join(current_path_parts)

        if isinstance(payload, str):
            if current_path and current_path in path_map:
                return path_map[current_path]
            return text_map.get(payload, payload)
        if isinstance(payload, list):
            result: list[Any] = []
            for index, item in enumerate(payload):
                next_part = f"[{index}]"
                next_path_parts = current_path_parts + [next_part]
                result.append(self._translate_tree(item, text_map, path_map, next_path_parts))
            return result
        if isinstance(payload, dict):
            result: dict[Any, Any] = {}
            for key, value in payload.items():
                if current_path_parts:
                    next_path_parts = current_path_parts + [str(key)]
                else:
                    next_path_parts = [str(key)]
                result[key] = self._translate_tree(value, text_map, path_map, next_path_parts)
            return result
        return payload

    def _decode_text_content(self, raw_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "utf-16", "gb18030", "iso-8859-1", "cp1252"):
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
            except UnicodeError:
                continue
        return raw_bytes.decode("utf-8", errors="replace")

    def _build_translated_filename(self, filename: str, extension_override: str | None = None) -> str:
        path = Path(filename or "translated.txt")
        extension = extension_override or path.suffix or ".txt"
        return f"{path.stem or 'translated'}_translated{extension}"

    def _get_mime_type(self, extension: str) -> str:
        mime_map = {
            ".csv": "text/csv; charset=utf-8",
            ".dxf": "image/vnd.dxf",
            ".dwg": "image/vnd.dwg",
            ".idml": "application/vnd.adobe.indesign-idml-package",
            ".json": "application/json; charset=utf-8",
            ".md": "text/markdown; charset=utf-8",
            ".markdown": "text/markdown; charset=utf-8",
            ".mif": "application/octet-stream",
            ".po": "text/x-gettext-translation; charset=utf-8",
            ".pot": "text/x-gettext-translation; charset=utf-8",
            ".properties": "text/plain; charset=utf-8",
            ".sdlxliff": "application/octet-stream",
            ".srt": "text/plain; charset=utf-8",
            ".strings": "text/plain; charset=utf-8",
            ".svg": "image/svg+xml",
            ".txml": "application/octet-stream",
            ".txt": "text/plain; charset=utf-8",
            ".dat": "text/plain; charset=utf-8",
            ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".yaml": "application/yaml; charset=utf-8",
            ".yml": "application/yaml; charset=utf-8",
            ".zip": "application/zip",
        }
        guessed = mimetypes.guess_type(f"file{extension}")[0]
        return mime_map.get(extension, guessed or "application/octet-stream")

    def _to_optional_int(self, value: Any) -> int | None:
        if value is None or value == "":
            return None
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    def _collapse_cad_sentence_segments(
        self,
        segments: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """导出前将同一 CAD 实体拆出的句段按原顺序合回一项。"""
        collapsed: list[dict[str, Any]] = []
        groups: dict[str, list[dict[str, Any]]] = {}
        group_items: dict[str, dict[str, Any]] = {}

        for segment in segments:
            metadata = segment.get("metadata", {}) or {}
            group_id = str(metadata.get("cad_sentence_group_id") or "")
            if not metadata.get("cad_sentence_split") or not group_id:
                collapsed.append(segment)
                continue
            if group_id not in groups:
                groups[group_id] = []
                aggregate = dict(segment)
                aggregate["metadata"] = dict(metadata)
                group_items[group_id] = aggregate
                collapsed.append(aggregate)
            groups[group_id].append(segment)

        for group_id, items in groups.items():
            items.sort(key=lambda item: int(
                (item.get("metadata", {}) or {}).get("cad_sentence_index", 0)
            ))
            aggregate = group_items[group_id]
            metadata = aggregate["metadata"]
            joiner = str(metadata.get("cad_sentence_joiner", " "))
            sources = [str(item.get("source_text") or "").strip() for item in items]
            targets = [str(item.get("target_text") or "").strip() for item in items]
            aggregate["source_text"] = str(
                metadata.get("cad_parent_source_text") or joiner.join(sources)
            )
            aggregate["display_text"] = str(
                metadata.get("cad_parent_display_text") or aggregate["source_text"]
            )
            translation_complete = bool(targets) and all(targets)
            metadata["cad_sentence_translation_complete"] = translation_complete
            aggregate["target_text"] = (
                joiner.join(targets)
                if translation_complete
                else ""
            )
        return collapsed

    def _extract_mtext_split_info(
        self,
        segments: list[dict[str, Any]],
        translations: dict[str, str],
    ) -> list[dict[str, Any]]:
        """从句段中提取 MTEXT 拆段信息，用于导出时替换单一原 MTEXT 为多个独立 MTEXT。

        每段带自身 y 位置，翻译后各段固定在原 y 上，不会因某段行数变化影响其他段。
        """
        grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for segment in segments:
            metadata = segment.get("metadata", {}) or {}
            parent = str(metadata.get("mtext_split_parent") or "")
            if parent:
                grouped[parent].append(segment)

        out: list[dict[str, Any]] = []
        for parent, parent_segments in grouped.items():
            # 拆段回写会先清空整个父 MTEXT，因此必须确保每个非空源段都有
            # 显式译文；任一段为空就保留原实体，禁止只重建成功的部分。
            if any(
                str(segment.get("source_text") or "").strip()
                and not str(segment.get("target_text") or "").strip()
                for segment in parent_segments
            ):
                logger.warning(
                    "跳过不完整 MTEXT 拆段回写 parent=%s segments=%d",
                    parent,
                    len(parent_segments),
                )
                continue

            for segment in parent_segments:
                metadata = segment.get("metadata", {}) or {}
                source_text = str(segment.get("source_text") or "").strip()
                target_text = str(segment.get("target_text") or "").strip()
                if not target_text:
                    continue
                out.append({
                    "parent_handle": parent,
                    "layout_version": metadata.get("mtext_split_layout_version", 1),
                    "indices": metadata.get("mtext_split_indices", []),
                    "source_text": source_text,
                    "target_text": target_text,
                    "x": metadata.get("primary_x") or metadata.get("group_x", 0),
                    "y": metadata.get("primary_y") or metadata.get("group_y_top", 0),
                    "height": metadata.get("primary_height", 2.5),
                    "width": metadata.get("group_width", 0),
                    "layer": metadata.get("layer", "0"),
                    "scope": metadata.get("scope", ""),
                    "style": metadata.get("primary_style", ""),
                    "color": metadata.get("primary_color", 256),
                    "true_color": metadata.get("primary_true_color"),
                    "transparency": metadata.get("primary_transparency"),
                    # 高度预算：本段 y 到下一段 y 的距离，导出时用于缩字号防砸表格
                    "y_budget": metadata.get("mtext_split_y_budget"),
                })
        return out

    def _extract_merged_text_info(
        self, 
        segments: list[dict[str, Any]], 
        translations: dict[str, str],
    ) -> list[dict[str, Any]]:
        """从句段中提取合并文本信息，用于 DWG/DXF 导出时的处理。
        
        支持两种合并方式：
        1. 自动空间合并（dwg_enable_spatial_merge=True 时）
        2. 手动合并（用户在工作台手动合并句段）
        
        导出时：
        - 把合并后的译文写到第一个实体的位置
        - 清空被合并的其他实体
        
        Args:
            segments: 句段列表
            translations: 源文本 -> 目标文本的映射
            
        Returns:
            合并文本信息列表
        """
        merged_info_list: list[dict[str, Any]] = []
        
        logger.debug("_extract_merged_text_info: 处理 %d 个句段", len(segments))
        
        for seg in segments:
            metadata = seg.get('metadata', {}) or {}

            # 同一个原 MTEXT 若已进入拆段回写，就不能再按闭合框重建，
            # 否则两条路径会各创建一套译文并叠在一起。
            if metadata.get('mtext_split_parent'):
                continue
            
            is_merged = metadata.get('is_merged', False)
            is_table_cell = metadata.get('cad_table_cell', False)
            entity_type = str(metadata.get('entity_type', '') or '').upper()
            # 单个 TEXT/ATTRIB/ATTDEF 也是有宽高范围的 CAD 文本块。兼容旧任务：
            # 即使尚无 cad_text_block 标记，只要保存了完整组几何也纳入候选。
            is_single_text_block = bool(
                not is_merged
                and not is_table_cell
                and entity_type in {'TEXT', 'ATTRIB', 'ATTDEF'}
                and (
                    metadata.get('cad_text_block', False)
                    or (
                        metadata.get('merged_handles')
                        and metadata.get('group_width')
                        and metadata.get('group_height')
                    )
                )
            )
            if not (is_merged or is_table_cell or is_single_text_block):
                continue
            
            logger.info(
                "_extract_merged_text_info: 发现合并句段 segment_id=%s, metadata=%s",
                seg.get('segment_id', ''), metadata
            )
            
            merged_handles = metadata.get('merged_handles', [])
            if not merged_handles:
                primary_metadata_handle = metadata.get('primary_handle') or metadata.get('handle', '')
                merged_handles = [primary_metadata_handle] if primary_metadata_handle else []
            if not merged_handles:
                logger.warning("_extract_merged_text_info: 文本框缺少实体 handle")
                continue
            
            primary_handle = metadata.get('primary_handle') or metadata.get('handle', '')
            if not primary_handle:
                primary_handle = merged_handles[0] if merged_handles else ''
            
            if not primary_handle:
                logger.warning("_extract_merged_text_info: 无法获取 primary_handle")
                continue
            
            # 获取译文
            source_text = str(seg.get('source_text', '')).strip()
            target_text = str(seg.get('target_text', '')).strip()
            
            if not target_text and source_text:
                target_text = translations.get(source_text, '')
            
            if not target_text:
                logger.warning(
                    "_extract_merged_text_info: 合并句段无译文 source=%s...",
                    source_text[:50] if source_text else "(empty)"
                )
                continue
            
            merged_info = {
                'source_text': source_text,
                'source_layout_text': str(
                    seg.get('display_text')
                    or metadata.get('cad_parent_display_text')
                    or source_text
                ),
                'target_text': target_text,
                'primary_handle': primary_handle,
                'merged_handles': merged_handles,
                'primary_x': metadata.get('primary_x') or metadata.get('x', 0),
                'primary_y': metadata.get('primary_y') or metadata.get('y', 0),
                'primary_height': metadata.get('primary_height') or metadata.get('height', 2.5),
                'primary_style': metadata.get('primary_style', ''),
                'primary_color': metadata.get('primary_color', 256),
                'primary_true_color': metadata.get('primary_true_color'),
                'primary_transparency': metadata.get('primary_transparency'),
                'group_x': metadata.get('group_x') or metadata.get('primary_x') or metadata.get('x', 0),
                'group_y_top': metadata.get('group_y_top') or metadata.get('primary_y') or metadata.get('y', 0),
                'group_width': metadata.get('group_width', 0),
                'group_height': metadata.get('group_height', 0),
                'first_line_indent': metadata.get('group_first_line_indent', 0),
                'cad_table_cell': is_table_cell,
                'single_text_block': is_single_text_block,
                'scope': metadata.get('scope', ''),
                'layer': metadata.get('layer', '0'),
            }
            merged_info_list.append(merged_info)
            
            logger.info(
                "_extract_merged_text_info: 添加合并组 primary=%s, handles=%s",
                primary_handle, merged_handles
            )
        
        if merged_info_list:
            logger.info(
                "DWG 导出合并信息提取：%d 个句段 → %d 个合并组",
                len(segments), len(merged_info_list)
            )
        else:
            logger.debug("_extract_merged_text_info: 未找到任何合并句段")
        
        return merged_info_list


def get_export_options_for_file(filename: str) -> list[dict]:
    return MultiFormatExporter().get_available_exports(filename)


def export_file(
    export_type: str,
    segments: list[Any],
    filename: str,
    original_bytes: bytes | None = None,
    source_lang: str = "zh-CN",
    target_lang: str = "en-US",
) -> tuple[bytes, str, str]:
    exporter = MultiFormatExporter(source_lang, target_lang)
    return exporter.export(export_type, segments, filename, original_bytes)
