"""
导出服务模块 - 将翻译后的 Document AST 导出为各种格式

Requirements: 12.1, 12.2, 12.3, 12.4, 12.5, 12.6, 12.7
"""
import zipfile
from io import BytesIO
from typing import Dict, List, Tuple

from docx import Document

from app.services.adapters.exceptions import ExportError
from app.services.adapters.models import BlockNode, DocumentAST, NodeType


class ExportService:
    """导出服务
    
    支持将翻译后的 Document AST 导出为 DOCX、TXT 等格式。
    """

    def export_docx(
        self,
        ast: DocumentAST,
        translations: Dict[str, str],
    ) -> bytes:
        """导出为 DOCX 格式
        
        Args:
            ast: 文档抽象语法树
            translations: segment_id -> 译文 的映射
            
        Returns:
            bytes: DOCX 文件字节
            
        Raises:
            ExportError: 当导出失败时
        """
        try:
            doc = Document()
            
            for node in ast.nodes:
                self._export_node_to_docx(doc, node, translations)
            
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except Exception as e:
            raise ExportError("DOCX", str(e))

    def export_txt(
        self,
        ast: DocumentAST,
        translations: Dict[str, str],
    ) -> bytes:
        """导出为 TXT 格式
        
        Args:
            ast: 文档抽象语法树
            translations: segment_id -> 译文 的映射
            
        Returns:
            bytes: TXT 文件字节（UTF-8 编码）
            
        Raises:
            ExportError: 当导出失败时
        """
        try:
            paragraphs = []
            
            for node in ast.nodes:
                text = self._export_node_to_text(node, translations)
                if text:
                    paragraphs.append(text)
            
            content = "\n\n".join(paragraphs)
            return content.encode("utf-8")
        except Exception as e:
            raise ExportError("TXT", str(e))

    def export_bilingual(
        self,
        ast: DocumentAST,
        translations: Dict[str, str],
        format: str = "txt",
    ) -> bytes:
        """导出双语对照文档
        
        Args:
            ast: 文档抽象语法树
            translations: segment_id -> 译文 的映射
            format: 输出格式（"txt" 或 "docx"）
            
        Returns:
            bytes: 文件字节
            
        Raises:
            ExportError: 当导出失败时
        """
        if format == "docx":
            return self._export_bilingual_docx(ast, translations)
        else:
            return self._export_bilingual_txt(ast, translations)

    def export_batch(
        self,
        documents: List[Tuple[str, DocumentAST, Dict[str, str]]],
        format: str = "docx",
    ) -> bytes:
        """批量导出为 ZIP 压缩包
        
        Args:
            documents: (文件名, AST, 翻译映射) 元组列表
            format: 输出格式（"txt" 或 "docx"）
            
        Returns:
            bytes: ZIP 文件字节
            
        Raises:
            ExportError: 当导出失败时
        """
        try:
            buffer = BytesIO()
            
            with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for filename, ast, translations in documents:
                    if format == "docx":
                        content = self.export_docx(ast, translations)
                        ext = ".docx"
                    else:
                        content = self.export_txt(ast, translations)
                        ext = ".txt"
                    
                    # 确保文件名有正确的扩展名
                    if not filename.endswith(ext):
                        filename = filename.rsplit(".", 1)[0] + ext
                    
                    zf.writestr(filename, content)
            
            return buffer.getvalue()
        except ExportError:
            raise
        except Exception as e:
            raise ExportError("ZIP", str(e))

    def _export_node_to_docx(
        self,
        doc: Document,
        node: BlockNode,
        translations: Dict[str, str],
    ) -> None:
        """将节点导出到 DOCX 文档
        
        Args:
            doc: python-docx 文档对象
            node: 块级节点
            translations: 翻译映射
        """
        if node.node_type == NodeType.PARAGRAPH:
            text = self._get_translated_text(node, translations)
            if text:
                doc.add_paragraph(text)
        
        elif node.node_type == NodeType.HEADING:
            text = self._get_translated_text(node, translations)
            level = node.metadata.get("level", 1)
            if text:
                doc.add_heading(text, level=level)
        
        elif node.node_type == NodeType.TABLE:
            self._export_table_to_docx(doc, node, translations)
        
        # 递归处理子节点
        for child in node.children:
            self._export_node_to_docx(doc, child, translations)

    def _export_table_to_docx(
        self,
        doc: Document,
        table_node: BlockNode,
        translations: Dict[str, str],
    ) -> None:
        """将表格节点导出到 DOCX
        
        Args:
            doc: python-docx 文档对象
            table_node: 表格节点
            translations: 翻译映射
        """
        rows = table_node.metadata.get("rows", 0)
        cols = table_node.metadata.get("columns", 0)
        
        if rows == 0 or cols == 0:
            return
        
        table = doc.add_table(rows=rows, cols=cols)
        
        for row_idx, row_node in enumerate(table_node.children):
            if row_idx >= rows:
                break
            
            for col_idx, cell_node in enumerate(row_node.children):
                if col_idx >= cols:
                    break
                
                text = self._get_translated_text(cell_node, translations)
                table.rows[row_idx].cells[col_idx].text = text or ""

    def _export_node_to_text(
        self,
        node: BlockNode,
        translations: Dict[str, str],
    ) -> str:
        """将节点导出为文本
        
        Args:
            node: 块级节点
            translations: 翻译映射
            
        Returns:
            str: 文本内容
        """
        if node.text_content:
            return self._get_translated_text(node, translations)
        
        # 递归处理子节点
        texts = []
        for child in node.children:
            text = self._export_node_to_text(child, translations)
            if text:
                texts.append(text)
        
        return " ".join(texts)

    def _get_translated_text(
        self,
        node: BlockNode,
        translations: Dict[str, str],
    ) -> str:
        """获取节点的翻译文本
        
        如果没有翻译，返回原文。
        
        Args:
            node: 块级节点
            translations: 翻译映射
            
        Returns:
            str: 翻译后的文本或原文
        """
        # 简单实现：直接返回原文或翻译
        # 实际应用中需要根据 segment_id 查找翻译
        original = node.text_content or ""
        
        # 尝试在翻译映射中查找
        for segment_id, translation in translations.items():
            if original in segment_id or segment_id in original:
                return translation
        
        return original

    def _export_bilingual_txt(
        self,
        ast: DocumentAST,
        translations: Dict[str, str],
    ) -> bytes:
        """导出双语 TXT 文档
        
        Args:
            ast: 文档抽象语法树
            translations: 翻译映射
            
        Returns:
            bytes: TXT 文件字节
        """
        lines = []
        
        for node in ast.nodes:
            source = node.text_content or ""
            target = self._get_translated_text(node, translations)
            
            if source:
                lines.append(f"[源文] {source}")
                lines.append(f"[译文] {target}")
                lines.append("")
        
        content = "\n".join(lines)
        return content.encode("utf-8")

    def _export_bilingual_docx(
        self,
        ast: DocumentAST,
        translations: Dict[str, str],
    ) -> bytes:
        """导出双语 DOCX 文档
        
        Args:
            ast: 文档抽象语法树
            translations: 翻译映射
            
        Returns:
            bytes: DOCX 文件字节
        """
        doc = Document()
        
        for node in ast.nodes:
            source = node.text_content or ""
            target = self._get_translated_text(node, translations)
            
            if source:
                # 创建双语表格
                table = doc.add_table(rows=1, cols=2)
                table.rows[0].cells[0].text = source
                table.rows[0].cells[1].text = target
                doc.add_paragraph()  # 添加空行分隔
        
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
