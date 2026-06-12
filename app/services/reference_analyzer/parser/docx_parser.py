"""Word 文档解析 - 支持表格提取、句子级分割和修订格式"""

import os
import re
from typing import Optional, List
from .base import BaseParser, Document, Paragraph, Table


class DocxParser(BaseParser):

    @staticmethod
    def supported_extensions():
        return [".docx", ".doc"]

    def parse(self, file_path: str) -> Document:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".doc":
            return self._parse_doc(file_path)
        else:
            return self._parse_docx(file_path)

    def _parse_docx(self, file_path: str) -> Document:
        """解析 .docx 文件，支持修订格式"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = []
        tables = []

        # 解析段落 - 使用完整的 XML 解析来处理修订格式
        for para in doc.paragraphs:
            # 使用增强方法提取完整文本（包括修订内容）
            text = self._extract_full_paragraph_text(para)
            if text:
                paragraphs.append(Paragraph(
                    text=text,
                    style=para.style.name if para.style else None,
                ))

        # 解析表格
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # 表格单元格也需要处理修订格式
                    cell_text = self._extract_full_cell_text(cell)
                    row_data.append(cell_text)
                rows_data.append(row_data)

            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []
            tables.append(Table(headers=headers, rows=data_rows))
            
            # 将表格内容也添加到段落中
            for row in rows_data:
                for cell_text in row:
                    if cell_text:
                        paragraphs.append(Paragraph(text=cell_text))

        raw_text = "\n".join(p.text for p in paragraphs)

        return Document(
            paragraphs=paragraphs,
            tables=tables,
            filename=os.path.basename(file_path),
            raw_text=raw_text,
        )

    def _extract_full_paragraph_text(self, para) -> str:
        """从段落中提取完整文本，包括修订格式中的内容
        
        Word 修订格式：
        - <w:ins> 标签包含插入的文本（应该保留）
        - <w:del> 标签包含删除的文本（应该忽略）
        - 默认的 para.text 可能不完整
        
        策略：遍历所有 run，提取接受修订后的完整文本
        """
        # 方法1: 先尝试遍历所有 runs
        # python-docx 的 runs 包含段落中的所有文本块
        text_from_runs = []
        try:
            for run in para.runs:
                if run.text:
                    text_from_runs.append(run.text)
        except Exception:
            pass
        
        text_v1 = ''.join(text_from_runs).strip()
        
        # 方法2: 使用 XML 深度遍历提取所有 <w:t> 文本
        text_v2 = self._extract_text_from_xml(para._element)
        
        # 方法3: 默认的 para.text
        text_v3 = para.text.strip() if para.text else ''
        
        # 选择最长的结果（通常最完整）
        candidates = [text_v1, text_v2, text_v3]
        best_text = max(candidates, key=len)
        
        # 调试输出（如果各方法结果不一致）
        if len(set(len(t) for t in candidates if t)) > 1:
            print(f"[DocxParser] 文本提取差异:")
            print(f"  runs方式: {len(text_v1)} 字符 - {text_v1[:50]}...")
            print(f"  XML方式:  {len(text_v2)} 字符 - {text_v2[:50]}...")
            print(f"  默认方式: {len(text_v3)} 字符 - {text_v3[:50]}...")
            print(f"  选择: {len(best_text)} 字符")
        
        return best_text

    def _extract_text_from_xml(self, element) -> str:
        """从 XML 元素中递归提取所有文本，跳过删除的内容"""
        text_parts = []
        
        # 遍历所有子元素
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            # 跳过删除标记内的内容
            if tag == 'del':
                continue
            
            # 文本节点
            if tag == 't':
                if child.text:
                    text_parts.append(child.text)
            else:
                # 递归处理子元素
                text_parts.append(self._extract_text_from_xml(child))
        
        return ''.join(text_parts)

    def _extract_full_cell_text(self, cell) -> str:
        """从表格单元格中提取完整文本"""
        text_parts = []
        
        for para in cell.paragraphs:
            para_text = self._extract_full_paragraph_text(para)
            if para_text:
                text_parts.append(para_text)
        
        return ' '.join(text_parts).strip()

    def _parse_doc(self, file_path: str) -> Document:
        """解析 .doc 文件 - 优先转换为 .docx 保留段落结构"""
        import subprocess
        import tempfile
        import platform
        
        # 方案1: 使用 LibreOffice 转换为 docx
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                soffice_paths = [
                    'soffice',
                    'soffice.exe',
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                    'libreoffice',
                ]
                
                for cmd in soffice_paths:
                    try:
                        result = subprocess.run(
                            [cmd, '--headless', '--convert-to', 'docx', '--outdir', tmpdir, file_path],
                            capture_output=True,
                            timeout=60,
                        )
                        if result.returncode == 0:
                            base_name = os.path.splitext(os.path.basename(file_path))[0]
                            docx_path = os.path.join(tmpdir, base_name + '.docx')
                            if os.path.exists(docx_path):
                                print(f"[DocParser] LibreOffice 转换 .doc -> .docx 成功")
                                return self._parse_docx(docx_path)
                    except FileNotFoundError:
                        continue
                    except Exception:
                        continue
        except Exception as e:
            print(f"[DocParser] LibreOffice 转换失败: {e}")
        
        # 方案2: Windows 使用 win32com
        if platform.system() == 'Windows':
            try:
                import win32com.client
                import pythoncom
                
                pythoncom.CoInitialize()
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    
                    paragraphs = []
                    for i in range(1, doc.Paragraphs.Count + 1):
                        text = doc.Paragraphs(i).Range.Text.strip()
                        if text:
                            paragraphs.append(Paragraph(text=text))
                    
                    raw_text = doc.Content.Text
                    doc.Close(False)
                    word.Quit()
                    
                    print(f"[DocParser] Word COM 解析 .doc 成功")
                    return Document(
                        paragraphs=paragraphs,
                        tables=[],
                        filename=os.path.basename(file_path),
                        raw_text=raw_text,
                    )
                finally:
                    pythoncom.CoUninitialize()
            except ImportError:
                print(f"[DocParser] win32com 未安装")
            except Exception as e:
                print(f"[DocParser] Word COM 失败: {e}")
        
        # 方案3: 备用 - 转为纯文本
        raw_text = ""
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                for cmd in ['soffice', r'C:\Program Files\LibreOffice\program\soffice.exe', 'libreoffice']:
                    try:
                        result = subprocess.run(
                            [cmd, '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, file_path],
                            capture_output=True,
                            timeout=60,
                        )
                        if result.returncode == 0:
                            base_name = os.path.splitext(os.path.basename(file_path))[0]
                            txt_path = os.path.join(tmpdir, base_name + '.txt')
                            if os.path.exists(txt_path):
                                with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    raw_text = f.read()
                                print(f"[DocParser] LibreOffice 转换 .doc -> .txt (备用)")
                                break
                    except:
                        continue
        except Exception as e:
            print(f"[DocParser] 备用方案失败: {e}")
        
        # 构建结果
        paragraphs = []
        if raw_text:
            for para_text in raw_text.split('\n'):
                para_text = para_text.strip()
                if para_text:
                    paragraphs.append(Paragraph(text=para_text))
        
        return Document(
            paragraphs=paragraphs,
            tables=[],
            filename=os.path.basename(file_path),
            raw_text=raw_text,
        )
