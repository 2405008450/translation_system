# ══════════════════════════════════════════════
# 核心渲染器 V2
# ══════════════════════════════════════════════
"""
DOCX 渲染器 V2：
- 精确还原 rFonts 四槽位
- 通过 ZIP 级别操作还原 Chart/Image 等嵌入资源
"""
import shutil
import os
import re
from pathlib import Path
from typing import Optional, Any
from zipfile import ZipFile, ZIP_DEFLATED
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from lxml import etree

from style_main import TextSegment, ParagraphStyle, _xml_str_to_elem, EmbeddedObject, RunStyle


class DocxRendererV2:
    ALIGNMENT_MAP: dict[str, int] = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }

    def __init__(self, source_path: Optional[str | Path] = None) -> None:
        """
        Args:
            source_path: 源 DOCX 路径。如果提供，将从中复制嵌入资源（chart/image等）
        """
        self.source_path = Path(source_path) if source_path else None
        self.document = Document()

        # 清除默认空段落
        for p in self.document.paragraphs:
            p._element.getparent().remove(p._element)

        # 如果有源文件，准备资源复制
        self._source_zip_entries: dict[str, bytes] = {}
        self._source_rels_xml: Optional[bytes] = None
        self._source_content_types: Optional[bytes] = None
        if self.source_path and self.source_path.exists():
            self._load_source_resources()

    def _load_source_resources(self) -> None:
        """预加载源文件中的所有 ZIP 条目"""
        try:
            with ZipFile(str(self.source_path), "r") as zf:
                for name in zf.namelist():
                    self._source_zip_entries[name] = zf.read(name)
        except Exception as e:
            print(f"[WARN] 加载源文件资源失败: {e}")

    def render(
        self,
        segments: list[TextSegment],
        output_path: str | Path,
        adjuster: Any = None,
    ) -> None:
        """渲染并保存 DOCX"""
        self._adjuster = adjuster
        output_path = Path(output_path)

        body_segs = [s for s in segments if s.source in ("正文", "表格")]
        header_segs = [s for s in segments if s.source == "页眉"]
        footer_segs = [s for s in segments if s.source == "页脚"]

        # 渲染正文
        self._render_body(body_segs)

        # 页眉页脚通过 ZIP 级别直接从源文档注入，不使用 python-docx API

        # 处理文档末尾的全局分节属性（如分栏、页面设置）
        # 注意：必须在页眉页脚渲染之后执行，避免 sectPr 替换导致 section 引用失效
        if body_segs and body_segs[-1].computed_para_style and body_segs[-1].computed_para_style.section_xml:
            try:
                last_sect_xml = body_segs[-1].computed_para_style.section_xml
                new_sect_elem = parse_xml(last_sect_xml)
                body = self.document._body._body
                old_sect = body.find(qn("w:sectPr"))
                if old_sect is not None:
                    # 用 replace 保持节点位置，避免破坏 section 对象的内部引用
                    old_sect.getparent().replace(old_sect, new_sect_elem)
                else:
                    body.append(new_sect_elem)
            except Exception as e:
                print(f"[WARN] 还原全局分节属性失败: {e}")

        # 先保存基础 DOCX
        temp_path = output_path.with_suffix(".tmp.docx")
        try:
            self.document.save(str(temp_path))

            # 如果有嵌入对象，需要在 ZIP 级别注入资源
            has_embedded = any(s.is_embedded_object for s in segments)
            if has_embedded and self._source_zip_entries:
                self._inject_embedded_resources(temp_path, output_path, segments)
            else:
                shutil.move(str(temp_path), str(output_path))

            # 注入 updateFields 设置，确保页码和目录在打开时自动更新
            self._inject_update_fields(output_path)

            print(f"\n[OK] 文件已保存至: {output_path.resolve()}")
        except Exception as e:
            raise RuntimeError(f"保存失败: {e}") from e
        finally:
            if temp_path.exists():
                try:
                    temp_path.unlink()
                except OSError:
                    pass

    # ──────────────────────────────────────
    # 正文渲染
    # ──────────────────────────────────────

    def _render_body(self, segments: list[TextSegment]) -> None:
        groups = self._group_segments(segments)
        for g in groups:
            match g["type"]:
                case "paragraph":
                    self._render_para_group(g["segments"], self.document)
                case "table":
                    self._render_table(g["segments"])
                case "sdt":
                    self._render_sdt_block(g["segments"][0])

    def _group_segments(self, segments: list[TextSegment]) -> list[dict]:
        groups: list[dict] = []
        i = 0
        while i < len(segments):
            seg = segments[i]
            if (
                seg.is_embedded_object
                and seg.embedded_object
                and seg.embedded_object.object_type == "sdt"
            ):
                groups.append({"type": "sdt", "segments": [seg]})
                i += 1
            elif seg.source == "正文":
                pid = seg.paragraph_id
                runs: list[TextSegment] = [seg]
                j = i + 1
                while j < len(segments) and segments[j].source == "正文" and segments[j].paragraph_id == pid:
                    runs.append(segments[j])
                    j += 1
                groups.append({"type": "paragraph", "segments": runs})
                i = j
            elif seg.source == "表格" and seg.table_meta:
                tid = seg.table_meta.table_id
                tsegs: list[TextSegment] = [seg]
                j = i + 1
                while (
                    j < len(segments)
                    and segments[j].source == "表格"
                    and segments[j].table_meta
                    and segments[j].table_meta.table_id == tid
                ):
                    tsegs.append(segments[j])
                    j += 1
                groups.append({"type": "table", "segments": tsegs})
                i = j
            else:
                groups.append({"type": "paragraph", "segments": [seg]})
                i += 1
        return groups

    # ──────────────────────────────────────
    # 段落渲染
    # ──────────────────────────────────────

    def _render_para_group(self, segments: list[TextSegment], container: Any) -> None:
        if not segments:
            return

        first = segments[0]

        if first.raw_paragraph_xml and self._para_has_complex_elements(first.raw_paragraph_xml):
            try:
                para_elem = _xml_str_to_elem(first.raw_paragraph_xml)
                if hasattr(container, "_body"):
                    body_elem = container._body._body
                    sect_pr = body_elem.find(qn("w:sectPr"))
                    if sect_pr is not None:
                        sect_pr.addprevious(para_elem)
                    else:
                        body_elem.append(para_elem)
                else:
                    container._element.append(para_elem)
                return
            except Exception as e:
                print(f"[WARN] 原始段落 XML 插入失败，回退到逐 run 渲染: {e}")

        para_style_name = first.paragraph_style_name
        paragraph = container.add_paragraph()

        if para_style_name:
            try:
                paragraph.style = para_style_name
            except Exception:
                try:
                    ppr = paragraph._element.get_or_add_pPr()
                    old_ps = ppr.find(qn("w:pStyle"))
                    if old_ps is not None:
                        ppr.remove(old_ps)
                    ps_elem = OxmlElement("w:pStyle")
                    ps_elem.set(qn("w:val"), para_style_name)
                    ppr.insert(0, ps_elem)
                except Exception:
                    pass

        if first.computed_para_style:
            self._apply_para_style(paragraph, first.computed_para_style)

        has_para_style = bool(first.paragraph_style_name)
        for seg in segments:
            if seg.is_embedded_object and seg.embedded_object:
                self._insert_embedded_run(paragraph, seg.embedded_object)
            elif seg.text:
                run = paragraph.add_run(seg.text)
                run_style = (
                    seg.direct_run_style if has_para_style and seg.direct_run_style is not None
                    else seg.computed_run_style
                )
                if run_style:
                    self._apply_run_style_precise(run, run_style)

    @staticmethod
    def _para_has_complex_elements(raw_xml: str) -> bool:
        return (
            "w:bookmarkStart" in raw_xml
            or "w:bookmarkEnd" in raw_xml
            or "w:hyperlink" in raw_xml
            or "w:sectPr" in raw_xml
            or "m:oMath" in raw_xml
        )

    @staticmethod
    def _strip_font_from_run_style(rs: RunStyle) -> RunStyle:
        """返回一个去掉字体/字号字段的 RunStyle 副本，保留其他格式（加粗/斜体等）"""
        import copy
        stripped = copy.copy(rs)
        stripped.font_ascii = None
        stripped.font_east_asia = None
        stripped.font_hAnsi = None
        stripped.font_cs = None
        stripped.font_size = None
        stripped.font_size_cs = None
        return stripped

    def _apply_para_style(self, para: Any, ps: ParagraphStyle) -> None:
        pf = para.paragraph_format

        if ps.alignment and ps.alignment in self.ALIGNMENT_MAP:
            pf.alignment = self.ALIGNMENT_MAP[ps.alignment]

        if ps.line_spacing is not None:
            if ps.line_spacing_rule in ("auto", None):
                pf.line_spacing = ps.line_spacing
            else:
                pf.line_spacing = Pt(ps.line_spacing)

        if ps.space_before is not None:
            pf.space_before = Pt(ps.space_before)
        if ps.space_after is not None:
            pf.space_after = Pt(ps.space_after)

        if ps.indent_left is not None:
            pf.left_indent = Pt(ps.indent_left)
        if ps.indent_right is not None:
            pf.right_indent = Pt(ps.indent_right)
        if ps.indent_first_line is not None and ps.indent_first_line > 0:
            pf.first_line_indent = Pt(ps.indent_first_line)
        if ps.indent_hanging is not None and ps.indent_hanging > 0:
            pf.first_line_indent = Pt(-ps.indent_hanging)

        if ps.keep_next is True:
            pf.keep_with_next = True
        if ps.keep_lines is True:
            pf.keep_together = True
        if ps.page_break_before is True:
            pf.page_break_before = True

        if ps.num_id is not None and ps.num_id_is_direct:
            ppr = para._element.get_or_add_pPr()
            old_numpr = ppr.find(qn("w:numPr"))
            if old_numpr is not None:
                ppr.remove(old_numpr)
            numpr_elem = OxmlElement("w:numPr")
            ilvl_elem = OxmlElement("w:ilvl")
            ilvl_elem.set(qn("w:val"), str(ps.num_ilvl if ps.num_ilvl is not None else 0))
            numid_elem = OxmlElement("w:numId")
            numid_elem.set(qn("w:val"), str(ps.num_id))
            numpr_elem.append(ilvl_elem)
            numpr_elem.append(numid_elem)
            ppr.append(numpr_elem)

        # 还原分节符/分栏配置 (Section Break / Columns)
        if ps.section_xml:
            try:
                sect_elem = parse_xml(ps.section_xml)
                ppr = para._element.get_or_add_pPr()
                # 如果已存在 sectPr，先移除旧的
                old_sect = ppr.find(qn("w:sectPr"))
                if old_sect is not None:
                    ppr.remove(old_sect)
                ppr.append(sect_elem)
            except Exception as e:
                print(f"[WARN] 还原分节符失败: {e}")

    # ──────────────────────────────────────
    # Run 精确样式还原（核心改进）
    # ──────────────────────────────────────

    def _apply_run_style_precise(self, run: Any, rs: RunStyle) -> None:
        """精确还原 Run 样式——直接操作 XML 确保四槽位字体正确"""
        run_elem = run._element

        # 确保 rPr 存在
        rpr = run_elem.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_elem.insert(0, rpr)

        # ── 字体：精确设置四槽位 ──
        if any([rs.font_ascii, rs.font_east_asia, rs.font_hAnsi, rs.font_cs]):
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            if rs.font_ascii:
                rfonts.set(qn("w:ascii"), rs.font_ascii)
            if rs.font_east_asia:
                rfonts.set(qn("w:eastAsia"), rs.font_east_asia)
            if rs.font_hAnsi:
                rfonts.set(qn("w:hAnsi"), rs.font_hAnsi)
            if rs.font_cs:
                rfonts.set(qn("w:cs"), rs.font_cs)
            # 同时设置 python-docx 的 font.name 以保持一致
            if rs.font_ascii:
                run.font.name = rs.font_ascii

        # ── 字号：同时设置 sz 和 szCs ──
        if rs.font_size is not None:
            self._set_rpr_elem(rpr, "w:sz", str(int(rs.font_size * 2)))
        if rs.font_size_cs is not None:
            self._set_rpr_elem(rpr, "w:szCs", str(int(rs.font_size_cs * 2)))

        # ── 加粗 ──
        if rs.bold is not None:
            self._set_toggle(rpr, "w:b", rs.bold)
        if rs.bold_cs is not None:
            self._set_toggle(rpr, "w:bCs", rs.bold_cs)

        # ── 斜体 ──
        if rs.italic is not None:
            self._set_toggle(rpr, "w:i", rs.italic)
        if rs.italic_cs is not None:
            self._set_toggle(rpr, "w:iCs", rs.italic_cs)

        # ── 下划线 ──
        if rs.underline and rs.underline != "none":
            self._set_rpr_elem(rpr, "w:u", rs.underline)

        # ── 删除线 ──
        if rs.strike is not None:
            self._set_toggle(rpr, "w:strike", rs.strike)
        if rs.dstrike is not None:
            self._set_toggle(rpr, "w:dstrike", rs.dstrike)

        # ── 颜色 ──
        if rs.color:
            hex_c = rs.color.lstrip("#")
            self._set_rpr_elem(rpr, "w:color", hex_c)

        # ── 高亮 ──
        if rs.highlight:
            self._set_rpr_elem(rpr, "w:highlight", rs.highlight)

        # ── 底纹 ──
        if rs.shading:
            hex_s = rs.shading.lstrip("#")
            shd = rpr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                rpr.append(shd)
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_s)

        # ── 小型大写 / 全部大写 ──
        if rs.small_caps is not None:
            self._set_toggle(rpr, "w:smallCaps", rs.small_caps)
        if rs.all_caps is not None:
            self._set_toggle(rpr, "w:caps", rs.all_caps)

        # ── 上下标 ──
        if rs.vertical_align and rs.vertical_align != "baseline":
            self._set_rpr_elem(rpr, "w:vertAlign", rs.vertical_align)

        # ── 字符间距 ──
        if rs.spacing is not None:
            sp = rpr.find(qn("w:spacing"))
            if sp is None:
                sp = OxmlElement("w:spacing")
                rpr.append(sp)
            sp.set(qn("w:val"), str(int(rs.spacing * 20)))

        # ── Issue #8: 字符缩放 w:w ──
        if rs.char_scale is not None:
            self._set_rpr_elem(rpr, "w:w", str(rs.char_scale))

        # ── Issue #8: fitText w:fitText ──
        if rs.fit_text is not None:
            ft = rpr.find(qn("w:fitText"))
            if ft is None:
                ft = OxmlElement("w:fitText")
                rpr.append(ft)
            ft.set(qn("w:val"), str(rs.fit_text))

        # ── 位置偏移 ──
        if rs.position is not None:
            self._set_rpr_elem(rpr, "w:position", str(int(rs.position * 2)))

        # ── 字距调整 ──
        if rs.kern is not None:
            self._set_rpr_elem(rpr, "w:kern", str(int(rs.kern * 2)))

        # ── 空心/阴影/阳文/阴文/隐藏 ──
        if rs.outline is not None:
            self._set_toggle(rpr, "w:outline", rs.outline)
        if rs.shadow is not None:
            self._set_toggle(rpr, "w:shadow", rs.shadow)
        if rs.emboss is not None:
            self._set_toggle(rpr, "w:emboss", rs.emboss)
        if rs.imprint is not None:
            self._set_toggle(rpr, "w:imprint", rs.imprint)
        if rs.vanish is not None:
            self._set_toggle(rpr, "w:vanish", rs.vanish)

        # ── 语言 ──
        if any([rs.lang_val, rs.lang_east_asia, rs.lang_bidi]):
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            if rs.lang_val:
                lang.set(qn("w:val"), rs.lang_val)
            if rs.lang_east_asia:
                lang.set(qn("w:eastAsia"), rs.lang_east_asia)
            if rs.lang_bidi:
                lang.set(qn("w:bidi"), rs.lang_bidi)

    @staticmethod
    def _set_rpr_elem(rpr: etree._Element, tag: str, val: str) -> None:
        """设置 rPr 子元素的 w:val 属性"""
        elem = rpr.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            rpr.append(elem)
        elem.set(qn("w:val"), val)

    @staticmethod
    def _set_toggle(rpr: etree._Element, tag: str, value: bool) -> None:
        """设置 toggle 类型属性"""
        elem = rpr.find(qn(tag))
        if value:
            if elem is None:
                elem = OxmlElement(tag)
                rpr.append(elem)
            # <w:b/> 等价于 <w:b w:val="true"/>，不需要显式设置 val
            # 但为了安全，移除可能存在的 val="0"
            if elem.get(qn("w:val")):
                del elem.attrib[qn("w:val")]
        else:
            if elem is None:
                elem = OxmlElement(tag)
                rpr.append(elem)
            elem.set(qn("w:val"), "0")

    # ──────────────────────────────────────
    # 嵌入对象插入
    # ──────────────────────────────────────

    def _insert_embedded_run(self, paragraph: Any, emb: EmbeddedObject) -> None:
        """将嵌入对象的原始 XML Run 插入段落"""
        try:
            elem = _xml_str_to_elem(emb.xml_snippet)
            paragraph._element.append(elem)
        except Exception as e:
            # 回退：插入占位符
            run = paragraph.add_run(f"[{emb.object_type}]")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            print(f"[WARN] 嵌入对象插入失败 ({emb.object_type}): {e}")

    # ──────────────────────────────────────
    # SDT（目录/结构化文档标签）渲染
    # ──────────────────────────────────────

    def _render_sdt_block(self, seg: "TextSegment") -> None:
        if not seg.embedded_object:
            return
        try:
            sdt_elem = _xml_str_to_elem(seg.embedded_object.xml_snippet)
            body = self.document._body._body
            sect_pr = body.find(qn("w:sectPr"))
            if sect_pr is not None:
                sect_pr.addprevious(sdt_elem)
            else:
                body.append(sdt_elem)
        except Exception as e:
            print(f"[WARN] SDT 块插入失败: {e}")

    # ──────────────────────────────────────
    # 表格渲染
    # ──────────────────────────────────────

    def _render_table(self, segments: list[TextSegment]) -> None:
        if not segments:
            return
        first_meta = segments[0].table_meta
        if not first_meta:
            return

        if first_meta.table_xml:
            try:
                table_xml = first_meta.table_xml
                if self._adjuster is not None and hasattr(self._adjuster, "apply_to_table_xml"):
                    # 优先用 "Table Grid" 或 defaults 规则处理表格
                    tbl_rule = (
                        self._adjuster._style_rules.get("Table Grid")
                        or self._adjuster._style_rules.get("表格网格")
                        or self._adjuster._defaults_rule
                    )
                    table_xml = self._adjuster.apply_to_table_xml(table_xml, tbl_rule)
                elif (
                    self._adjuster is not None
                    and hasattr(self._adjuster, "strip_font_from_table_xml")
                    and (self._adjuster._style_rules or self._adjuster._defaults_rule is not None)
                ):
                    table_xml = self._adjuster.strip_font_from_table_xml(table_xml)
                tbl_elem = _xml_str_to_elem(table_xml)
                body = self.document._body._body
                sect_pr = body.find(qn("w:sectPr"))
                if sect_pr is not None:
                    sect_pr.addprevious(tbl_elem)
                else:
                    body.append(tbl_elem)
                return
            except Exception as e:
                print(f"[WARN] 原始表格 XML 插入失败，回退到重建模式: {e}")

        total_rows = first_meta.total_rows
        total_cols = first_meta.total_cols
        if total_rows == 0 or total_cols == 0:
            return

        table = self.document.add_table(rows=total_rows, cols=total_cols)
        table.style = "Table Grid"

        if first_meta.table_pr_xml:
            try:
                tbl_pr_elem = parse_xml(first_meta.table_pr_xml)
                tbl_elem = table._element
                old_tbl_pr = tbl_elem.find(qn("w:tblPr"))
                if old_tbl_pr is not None:
                    tbl_elem.remove(old_tbl_pr)
                tbl_elem.insert(0, tbl_pr_elem)
            except Exception as e:
                print(f"[WARN] 还原表格属性失败: {e}")

        cell_data: dict[tuple[int, int], list[list[TextSegment]]] = {}
        for seg in segments:
            if not seg.table_meta:
                continue
            key = (seg.table_meta.row, seg.table_meta.col)
            if key not in cell_data:
                cell_data[key] = []
            if cell_data[key] and cell_data[key][-1][0].paragraph_id == seg.paragraph_id:
                cell_data[key][-1].append(seg)
            else:
                cell_data[key].append([seg])

        table_cells = table._cells
        for (row, col), para_groups in cell_data.items():
            if row >= total_rows or col >= total_cols:
                continue
            cell = table_cells[row * total_cols + col]
            meta = para_groups[0][0].table_meta

            if meta and meta.cell_pr_xml:
                try:
                    tc_pr_elem = parse_xml(meta.cell_pr_xml)
                    tc_elem = cell._element
                    old_tc_pr = tc_elem.find(qn("w:tcPr"))
                    if old_tc_pr is not None:
                        tc_elem.remove(old_tc_pr)
                    tc_elem.insert(0, tc_pr_elem)
                except Exception as e:
                    print(f"[WARN] 还原单元格属性失败 (R{row}C{col}): {e}")

            while len(cell.paragraphs) > 1:
                last_p = cell.paragraphs[-1]._element
                last_p.getparent().remove(last_p)

            for pg_idx, para_runs in enumerate(para_groups):
                if pg_idx == 0:
                    paragraph = cell.paragraphs[0]
                else:
                    paragraph = cell.add_paragraph()

                if para_runs and para_runs[0].computed_para_style:
                    self._apply_para_style(paragraph, para_runs[0].computed_para_style)

                for seg in para_runs:
                    if seg.is_embedded_object and seg.embedded_object:
                        self._insert_embedded_run(paragraph, seg.embedded_object)
                    elif seg.text:
                        run = paragraph.add_run(seg.text)
                        if seg.computed_run_style:
                            self._apply_run_style_precise(run, seg.computed_run_style)

    # ──────────────────────────────────────
    # 页眉页脚渲染
    # ──────────────────────────────────────

    def _render_header_footer(self, segments: list[TextSegment], is_header: bool) -> None:
        try:
            section = self.document.sections[0]
            container = section.header if is_header else section.footer

            try:
                container.is_linked_to_previous = False
            except Exception:
                sect_pr = self.document._body._body.find(qn("w:sectPr"))
                if sect_pr is not None:
                    ref_tag = qn("w:headerReference") if is_header else qn("w:footerReference")
                    for ref in list(sect_pr.findall(ref_tag)):
                        sect_pr.remove(ref)
                    new_ref = OxmlElement(ref_tag)
                    new_ref.set(qn("w:type"), "default")
                    sect_pr.append(new_ref)

            for p in list(container.paragraphs):
                p._element.getparent().remove(p._element)

            groups = self._group_by_para_id(segments)
            for group in groups:
                self._render_para_group(group, container)
        except Exception as e:
            print(f"[WARN] 页眉页脚渲染失败: {e}")

    @staticmethod
    def _group_by_para_id(segments: list[TextSegment]) -> list[list[TextSegment]]:
        groups: list[list[TextSegment]] = []
        cur_pid: Optional[int] = None
        for seg in segments:
            if seg.paragraph_id != cur_pid:
                groups.append([seg])
                cur_pid = seg.paragraph_id
            else:
                groups[-1].append(seg)
        return groups

    # ──────────────────────────────────────
    # ZIP 级别嵌入资源注入
    # ──────────────────────────────────────

    def _inject_embedded_resources(
        self,
        temp_docx: Path,
        output_path: Path,
        segments: list[TextSegment],
    ) -> None:
        """
        在 ZIP 级别将源文件中的 chart/image/ole 等资源注入到输出文件中。

        策略：
        1. 收集所有嵌入对象引用的 rId
        2. 从源文件的 word/_rels/document.xml.rels 中找到对应的 target
        3. 将 target 文件（chart xml、embedded xlsx、image 等）从源 ZIP 复制到输出 ZIP
        4. 将 relationship 条目添加到输出文件的 rels 中
        5. 更新 [Content_Types].xml
        """
        # 收集所有需要的 rId
        needed_rids: set[str] = set()
        for seg in segments:
            if seg.is_embedded_object and seg.embedded_object:
                needed_rids.update(seg.embedded_object.relationship_ids)

        if not needed_rids:
            shutil.move(str(temp_docx), str(output_path))
            return

        print(f"   [INFO] 需要注入 {len(needed_rids)} 个嵌入资源引用")

        rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

        # 构建全局 rId -> (type, target, base_dir, src_rels_file) 映射
        # 扫描源文件中所有 word/_rels/*.rels，覆盖 document/header/footer 等
        # rid_map: rId -> (rtype, target, base_dir, src_rels_file)
        rid_map: dict[str, tuple[str, str, str, str]] = {}
        for rels_entry, rels_data in self._source_zip_entries.items():
            if not rels_entry.startswith("word/_rels/") or not rels_entry.endswith(".rels"):
                continue
            base_dir = rels_entry.split("/_rels/")[0]  # "word"
            try:
                root = etree.fromstring(rels_data)
                for rel in root:
                    rid = rel.get("Id", "")
                    rtype = rel.get("Type", "")
                    target = rel.get("Target", "")
                    if rid and rid not in rid_map:
                        rid_map[rid] = (rtype, target, base_dir, rels_entry)
            except Exception:
                pass

        source_rels_path = "word/_rels/document.xml.rels"
        if source_rels_path not in self._source_zip_entries:
            print("   [WARN] 源文件缺少 document.xml.rels，跳过资源注入")
            shutil.move(str(temp_docx), str(output_path))
            return

        # 确定需要复制的文件，并按来源 rels 文件分组
        files_to_copy: dict[str, bytes] = {}  # zip_path -> bytes
        # rels_to_add_by_file: 输出 rels 文件路径 -> [(rId, type, target)]
        rels_to_add_by_file: dict[str, list[tuple[str, str, str]]] = {}

        for rid in needed_rids:
            if rid not in rid_map:
                print(f"   [WARN] rId {rid} 在所有 rels 中均未找到")
                continue
            rtype, target, base_dir, src_rels_file = rid_map[rid]

            # relationship 写回它原本所在的 rels 文件
            rels_to_add_by_file.setdefault(src_rels_file, []).append((rid, rtype, target))

            # 解析完整 ZIP 路径
            if target.startswith("/"):
                full_path = target.lstrip("/")
            else:
                full_path = f"{base_dir}/{target}"
            full_path = os.path.normpath(full_path).replace("\\", "/")

            if full_path in self._source_zip_entries:
                files_to_copy[full_path] = self._source_zip_entries[full_path]
            else:
                print(f"   [WARN] 资源文件不存在于源 ZIP: {full_path}")

            self._collect_dependent_files(full_path, files_to_copy)

        # 重建 ZIP：直接用源文件的 document.xml.rels 替换输出文件的
        # 原因：python-docx 生成的 rels 会占用相同的 rId（如 rId7=fontTable），
        # 与源文件图片的 rId 冲突，导致图片 relationship 被跳过。
        # 源文件的 rels 已包含所有必要引用，直接替换最可靠。
        src_doc_rels = self._source_zip_entries.get(source_rels_path)

        with ZipFile(str(temp_docx), "r") as src_zip:
            with ZipFile(str(output_path), "w", ZIP_DEFLATED) as dst_zip:
                for item in src_zip.namelist():
                    if item == source_rels_path and src_doc_rels:
                        # 直接用源文件的 document.xml.rels
                        dst_zip.writestr(item, src_doc_rels)
                    elif item == "[Content_Types].xml":
                        data = self._merge_content_types(src_zip.read(item), files_to_copy)
                        dst_zip.writestr(item, data)
                    else:
                        dst_zip.writestr(item, src_zip.read(item))

                existing = set(src_zip.namelist())

                # 如果某个非 document rels 文件在输出 docx 里不存在，从源文件补充
                for rels_file, new_rels in rels_to_add_by_file.items():
                    if rels_file == source_rels_path:
                        continue  # 已经用源文件替换了
                    if rels_file not in existing:
                        src_rels_data = self._source_zip_entries.get(rels_file)
                        if src_rels_data:
                            dst_zip.writestr(rels_file, src_rels_data)
                        else:
                            root_elem = etree.Element(f"{{{rels_ns}}}Relationships")
                            for rid, rtype, target in new_rels:
                                rel_elem = etree.SubElement(root_elem, f"{{{rels_ns}}}Relationship")
                                rel_elem.set("Id", rid)
                                rel_elem.set("Type", rtype)
                                rel_elem.set("Target", target)
                            data = etree.tostring(root_elem, xml_declaration=True, encoding="UTF-8", standalone=True)
                            dst_zip.writestr(rels_file, data)
                        print(f"   [INJECT] 新建 rels: {rels_file}")

                # 添加嵌入资源文件
                for zip_path, content in files_to_copy.items():
                    if zip_path not in existing:
                        dst_zip.writestr(zip_path, content)
                        print(f"   [INJECT] 注入: {zip_path}")

        print(f"   [OK] 嵌入资源注入完成")

    def _collect_dependent_files(self, file_path: str, files: dict[str, bytes]) -> None:
        """递归收集文件的依赖（子 rels、嵌入的 xlsx、图片等）"""
        # 查找对应的 rels 文件
        parts = file_path.rsplit("/", 1)
        if len(parts) == 2:
            rels_path = f"{parts[0]}/_rels/{parts[1]}.rels"
        else:
            rels_path = f"_rels/{parts[0]}.rels"

        if rels_path in self._source_zip_entries and rels_path not in files:
            rels_data = self._source_zip_entries[rels_path]
            files[rels_path] = rels_data

            # 解析子 rels 中引用的文件
            try:
                rels_root = etree.fromstring(rels_data)
                base_dir = parts[0] if len(parts) == 2 else ""
                for rel in rels_root:
                    target = rel.get("Target", "")
                    if not target or target.startswith("http"):
                        continue
                    if target.startswith("/"):
                        dep_path = target.lstrip("/")
                    elif base_dir:
                        # 处理相对路径中的 ..
                        dep_path = os.path.normpath(f"{base_dir}/{target}").replace("\\", "/")
                    else:
                        dep_path = target

                    if dep_path in self._source_zip_entries and dep_path not in files:
                        files[dep_path] = self._source_zip_entries[dep_path]
                        # 递归
                        self._collect_dependent_files(dep_path, files)
            except Exception:
                pass

        # 特殊处理：chart 可能嵌入 xlsx
        if "charts/" in file_path:
            # 检查同目录下的 style/colors XML
            base = file_path.rsplit(".", 1)[0]
            for suffix in ["style1.xml", "colors1.xml"]:
                # 尝试常见命名模式
                chart_num = re.search(r"chart(\d+)", file_path)
                if chart_num:
                    num = chart_num.group(1)
                    for pattern in [
                        f"word/charts/style{num}.xml",
                        f"word/charts/colors{num}.xml",
                    ]:
                        if pattern in self._source_zip_entries and pattern not in files:
                            files[pattern] = self._source_zip_entries[pattern]

    @staticmethod
    def _merge_rels(
        rels_data: bytes,
        new_rels: list[tuple[str, str, str]],
        rels_ns: str,
    ) -> bytes:
        """将新的 relationship 条目合并到 rels XML 中"""
        root = etree.fromstring(rels_data)
        existing_ids = {rel.get("Id") for rel in root}

        for rid, rtype, target in new_rels:
            if rid in existing_ids:
                continue
            rel_elem = etree.SubElement(root, f"{{{rels_ns}}}Relationship")
            rel_elem.set("Id", rid)
            rel_elem.set("Type", rtype)
            rel_elem.set("Target", target)

        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    def _inject_update_fields(self, docx_path: Path) -> None:
        """注入 updateFields、numbering.xml、header/footer"""
        settings_path = "word/settings.xml"
        numbering_path = "word/numbering.xml"
        doc_path = "word/document.xml"
        rels_path = "word/_rels/document.xml.rels"
        ct_path = "[Content_Types].xml"
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        pkg_rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
        numbering_rel_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
        )
        header_rel_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
        )
        footer_rel_type = (
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
        )

        source_numbering = self._source_zip_entries.get(numbering_path) if self._source_zip_entries else None
        source_styles = self._source_zip_entries.get("word/styles.xml") if self._source_zip_entries else None

        # 收集源文档的所有 header/footer 文件及其关系
        src_hf_files: dict[str, bytes] = {}   # 仅 header/footer XML 本身
        src_hf_extra: dict[str, bytes] = {}   # rels 和 media（不覆盖已注入版本）
        src_hf_rels: list[dict] = []
        if self._source_zip_entries:
            for key, data in self._source_zip_entries.items():
                fname = key.split("/")[-1]
                if (fname.startswith("header") or fname.startswith("footer")) and fname.endswith(".xml"):
                    src_hf_files[key] = data
                    # 收集页眉/页脚对应的 rels 文件及其引用的资源（图片等）
                    hf_rels_path = f"word/_rels/{fname}.rels"
                    if hf_rels_path in self._source_zip_entries:
                        src_hf_extra[hf_rels_path] = self._source_zip_entries[hf_rels_path]
                        # 解析 rels，收集引用的 media/image 等文件
                        try:
                            hf_rels_root = etree.fromstring(self._source_zip_entries[hf_rels_path])
                            for rel in hf_rels_root:
                                target = rel.get("Target", "")
                                if not target or target.startswith("http"):
                                    continue
                                if target.startswith("/"):
                                    dep_path = target.lstrip("/")
                                else:
                                    dep_path = os.path.normpath(f"word/{target}").replace("\\", "/")
                                if dep_path in self._source_zip_entries:
                                    src_hf_extra[dep_path] = self._source_zip_entries[dep_path]
                        except Exception as e:
                            print(f"[WARN] 解析页眉页脚 rels 失败 ({hf_rels_path}): {e}")

            src_rels_data = self._source_zip_entries.get(rels_path)
            if src_rels_data:
                try:
                    src_rels_root = etree.fromstring(src_rels_data)
                    for rel in src_rels_root:
                        rtype = rel.get("Type", "")
                        if rtype in (header_rel_type, footer_rel_type):
                            src_hf_rels.append({
                                "Id": rel.get("Id"),
                                "Type": rtype,
                                "Target": rel.get("Target"),
                            })
                except Exception as e:
                    print(f"[WARN] 解析源文档 rels 失败: {e}")

        tmp = docx_path.with_suffix(".uf.tmp.docx")
        try:
            with ZipFile(str(docx_path), "r") as src_zip:
                with ZipFile(str(tmp), "w", ZIP_DEFLATED) as dst_zip:
                    for item in src_zip.namelist():
                        # 跳过将被覆盖的文件
                        if source_numbering and item == numbering_path:
                            continue
                        if source_styles and item == "word/styles.xml":
                            continue
                        if src_hf_files and item in src_hf_files:
                            continue

                        data = src_zip.read(item)

                        if item == settings_path:
                            try:
                                root = etree.fromstring(data)
                                uf_tag = f"{{{w_ns}}}updateFields"
                                existing = root.find(uf_tag)
                                if existing is None:
                                    uf_elem = etree.SubElement(root, uf_tag)
                                    uf_elem.set(f"{{{w_ns}}}val", "true")
                                else:
                                    existing.set(f"{{{w_ns}}}val", "true")
                                data = etree.tostring(
                                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                                )
                            except Exception as e:
                                print(f"[WARN] 修改 settings.xml 失败: {e}")

                        elif item == rels_path:
                            try:
                                root = etree.fromstring(data)
                                existing_types = {rel.get("Type", "") for rel in root}
                                existing_ids = {rel.get("Id", "") for rel in root}

                                if source_numbering and numbering_rel_type not in existing_types:
                                    new_id = "rIdNumbering"
                                    c = 1
                                    while new_id in existing_ids:
                                        new_id = f"rIdNumbering{c}"; c += 1
                                    rel_elem = etree.SubElement(root, f"{{{pkg_rels_ns}}}Relationship")
                                    rel_elem.set("Id", new_id)
                                    rel_elem.set("Type", numbering_rel_type)
                                    rel_elem.set("Target", "numbering.xml")
                                    existing_ids.add(new_id)

                                for hf_rel in src_hf_rels:
                                    rid = hf_rel["Id"]
                                    if rid not in existing_ids:
                                        rel_elem = etree.SubElement(root, f"{{{pkg_rels_ns}}}Relationship")
                                        rel_elem.set("Id", rid)
                                        rel_elem.set("Type", hf_rel["Type"])
                                        rel_elem.set("Target", hf_rel["Target"])
                                        existing_ids.add(rid)

                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                print(f"[WARN] 更新 rels 失败: {e}")

                        elif item == ct_path:
                            try:
                                ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
                                root = etree.fromstring(data)
                                existing_parts = {ov.get("PartName", "") for ov in root.findall(f"{{{ct_ns}}}Override")}
                                existing_defaults = {d.get("Extension", "").lower() for d in root.findall(f"{{{ct_ns}}}Default")}

                                if source_numbering and "/word/numbering.xml" not in existing_parts:
                                    ov = etree.SubElement(root, f"{{{ct_ns}}}Override")
                                    ov.set("PartName", "/word/numbering.xml")
                                    ov.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml")

                                media_ct_map = {
                                    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                                    "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
                                    "emf": "image/x-emf", "wmf": "image/x-wmf",
                                    "svg": "image/svg+xml",
                                }
                                for hf_path_key in src_hf_files:
                                    part_name = "/" + hf_path_key
                                    fname_hf = hf_path_key.split("/")[-1]
                                    if fname_hf.startswith("header") and fname_hf.endswith(".xml"):
                                        if part_name not in existing_parts:
                                            ov = etree.SubElement(root, f"{{{ct_ns}}}Override")
                                            ov.set("PartName", part_name)
                                            ov.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml")
                                    elif fname_hf.startswith("footer") and fname_hf.endswith(".xml"):
                                        if part_name not in existing_parts:
                                            ov = etree.SubElement(root, f"{{{ct_ns}}}Override")
                                            ov.set("PartName", part_name)
                                            ov.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml")

                                # 注册 src_hf_extra 里 media 文件的扩展名
                                for extra_path in src_hf_extra:
                                    ext = extra_path.rsplit(".", 1)[-1].lower() if "." in extra_path else ""
                                    if ext in media_ct_map and ext not in existing_defaults:
                                        de = etree.SubElement(root, f"{{{ct_ns}}}Default")
                                        de.set("Extension", ext)
                                        de.set("ContentType", media_ct_map[ext])
                                        existing_defaults.add(ext)

                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                print(f"[WARN] 更新 Content_Types.xml 失败: {e}")

                        elif item == doc_path and src_hf_rels:
                            try:
                                root = etree.fromstring(data)
                                body_elem = root.find(f"{{{w_ns}}}body")
                                if body_elem is not None:
                                    sect_pr = body_elem.find(f"{{{w_ns}}}sectPr")
                                    if sect_pr is None:
                                        sect_pr = etree.SubElement(body_elem, f"{{{w_ns}}}sectPr")
                                    # 移除旧的 headerReference/footerReference
                                    for old in list(sect_pr.findall(f"{{{w_ns}}}headerReference")) + list(sect_pr.findall(f"{{{w_ns}}}footerReference")):
                                        sect_pr.remove(old)
                                    # 从源文档 sectPr 复制 headerReference/footerReference
                                    src_doc_data = self._source_zip_entries.get(doc_path)
                                    if src_doc_data:
                                        src_doc_root = etree.fromstring(src_doc_data)
                                        src_body = src_doc_root.find(f"{{{w_ns}}}body")
                                        if src_body is not None:
                                            src_sect = src_body.find(f"{{{w_ns}}}sectPr")
                                            if src_sect is not None:
                                                for ref in list(src_sect.findall(f"{{{w_ns}}}headerReference")) + list(src_sect.findall(f"{{{w_ns}}}footerReference")):
                                                    import copy as _copy
                                                    sect_pr.insert(0, _copy.deepcopy(ref))
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                            except Exception as e:
                                print(f"[WARN] 更新 document.xml sectPr 页眉页脚引用失败: {e}")

                        dst_zip.writestr(item, data)

                    if source_styles:
                        styles_data = source_styles
                        if self._adjuster is not None and hasattr(self._adjuster, "apply_to_styles_xml"):
                            styles_data = self._adjuster.apply_to_styles_xml(source_styles)
                        dst_zip.writestr("word/styles.xml", styles_data)
                        print(f"   [INJECT] 注入/覆盖: word/styles.xml")
                    if source_numbering:
                        dst_zip.writestr(numbering_path, source_numbering)
                        print(f"   [INJECT] 注入/覆盖: {numbering_path}")

                    for hf_path_key, hf_data in src_hf_files.items():
                        dst_zip.writestr(hf_path_key, hf_data)
                        print(f"   [INJECT] 注入页眉页脚: {hf_path_key}")

                    # rels 和 media 文件：只在输出 ZIP 里不存在时才写入（保留已注入的版本）
                    written = set(src_zip.namelist()) | set(src_hf_files.keys())
                    for extra_path, extra_data in src_hf_extra.items():
                        if extra_path not in written:
                            dst_zip.writestr(extra_path, extra_data)
                            print(f"   [INJECT] 注入页眉页脚资源: {extra_path}")

            os.replace(str(tmp), str(docx_path))
        except Exception as e:
            print(f"[WARN] 注入失败: {e}")
            if tmp.exists():
                tmp.unlink(missing_ok=True)

    @staticmethod
    def _merge_content_types(ct_data: bytes, new_files: dict[str, bytes]) -> bytes:
        """更新 [Content_Types].xml，添加新文件的内容类型"""
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        root = etree.fromstring(ct_data)

        # 已有的 Override 路径
        existing_overrides = set()
        for override in root.findall(f"{{{ct_ns}}}Override"):
            existing_overrides.add(override.get("PartName", ""))

        # 已有的 Default 扩展名
        existing_defaults = set()
        for default in root.findall(f"{{{ct_ns}}}Default"):
            existing_defaults.add(default.get("Extension", "").lower())

        # 内容类型映射
        content_type_map = {
            ".xml": "application/xml",
            ".rels": "application/vnd.openxmlformats-package.relationships+xml",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".tiff": "image/tiff",
            ".emf": "image/x-emf",
            ".wmf": "image/x-wmf",
            ".bin": "application/vnd.openxmlformats-officedocument.oleObject",
        }

        chart_ct = "application/vnd.openxmlformats-officedocument.drawingml.chart+xml"
        chart_style_ct = "application/vnd.ms-office.chartstyle+xml"
        chart_colors_ct = "application/vnd.ms-office.chartcolorstyle+xml"

        for file_path in new_files:
            part_name = f"/{file_path}" if not file_path.startswith("/") else file_path

            if part_name in existing_overrides:
                continue

            # 确定内容类型
            ext = Path(file_path).suffix.lower()
            if "chart" in file_path.lower() and ext == ".xml":
                if "style" in file_path.lower():
                    ct = chart_style_ct
                elif "color" in file_path.lower():
                    ct = chart_colors_ct
                else:
                    ct = chart_ct
            else:
                ct = content_type_map.get(ext)

            if ct and ext not in (".rels",):
                override = etree.SubElement(root, f"{{{ct_ns}}}Override")
                override.set("PartName", part_name)
                override.set("ContentType", ct)

            # 确保扩展名有 Default
            ext_no_dot = ext.lstrip(".")
            if ext_no_dot and ext_no_dot not in existing_defaults:
                default_ct = content_type_map.get(ext)
                if default_ct:
                    default_elem = etree.SubElement(root, f"{{{ct_ns}}}Default")
                    default_elem.set("Extension", ext_no_dot)
                    default_elem.set("ContentType", default_ct)
                    existing_defaults.add(ext_no_dot)

        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
