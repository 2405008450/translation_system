import copy
from pathlib import Path
from typing import Optional, Any
from zipfile import ZipFile, ZIP_DEFLATED
import json
from docx import Document
from docx.oxml.ns import qn, nsmap
from lxml import etree
from style_main import (
    RunStyle, ParagraphStyle, _get_toggle, _get_toggle_raw, _half_pt, _twips,
    _find_relationship_ids, _elem_to_xml_str, EmbeddedObject, ALL_NS,
    TextSegment, TableMeta, A_NS, _TOGGLE_FIELDS,
)

# ══════════════════════════════════════════════
# 增强版提取器
# ══════════════════════════════════════════════
"""
增强提取器 V2：
- 精确计算 Run 样式（四槽位字体）
- 识别并保存 Drawing/Chart/Image 等嵌入对象的原始 XML
- 记录表格行列坐标和段落 ID
"""
class EnhancedExtractorV2:
    def __init__(self, file_path: str | Path) -> None:
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")

        self.document = Document(str(self.file_path))
        self._style_cache: dict[str, etree._Element] = {}
        self._resolved_run_cache: dict[str, RunStyle] = {}
        self._resolved_para_cache: dict[str, ParagraphStyle] = {}
        self._theme_fonts = RunStyle()
        self._theme_colors: dict[str, str] = {}   # Issue #4: themeColor name -> hex
        self._doc_defaults_rpr = RunStyle()
        self._doc_defaults_ppr = ParagraphStyle()
        self._normal_style_id: Optional[str] = None
        self._para_id_counter = 0
        self._table_id_counter = 0

        # Issue #2: numbering.xml cache  numId -> {ilvl -> {pPr/rPr RunStyle/ParagraphStyle}}
        self._numbering_cache: dict[int, dict[int, dict]] = {}

        # Issue #3: table style resolved cache
        self._resolved_tbl_run_cache: dict[str, RunStyle] = {}
        self._resolved_tbl_para_cache: dict[str, ParagraphStyle] = {}

        # 收集源文件中所有 relationship
        self._source_rels: dict[str, dict] = {}
        self._collect_source_rels()

        self._build_style_cache()
        self._parse_theme()
        self._parse_doc_defaults()
        self._parse_numbering()   # Issue #2

    def _collect_source_rels(self) -> None:
        """收集源文档的所有 relationship 信息"""
        try:
            main_part = self.document.part
            for rel_id, rel in main_part.rels.items():
                self._source_rels[rel_id] = {
                    "reltype": rel.reltype,
                    "target_ref": rel.target_ref if hasattr(rel, "target_ref") else str(rel._target),
                    "is_external": rel.is_external,
                }
        except Exception:
            pass

    def _build_style_cache(self) -> None:
        try:
            styles_part = self.document.part._styles_part
            if styles_part is None:
                return
            for style_elem in styles_part._element.findall(qn("w:style")):
                sid = style_elem.get(qn("w:styleId"), "")
                if sid:
                    self._style_cache[sid] = style_elem
                    name_elem = style_elem.find(qn("w:name"))
                    if name_elem is not None:
                        name_val = name_elem.get(qn("w:val"), "")
                        if name_val.lower() in ("normal", "正文"):
                            self._normal_style_id = sid
        except Exception:
            pass

    def _parse_theme(self) -> None:
        try:
            with ZipFile(str(self.file_path), "r") as zf:
                if "word/theme/theme1.xml" not in zf.namelist():
                    return
                root = etree.fromstring(zf.read("word/theme/theme1.xml"))
            ns = {"a": A_NS}

            # ── 字体 ──
            minor = root.find(".//a:themeElements/a:fontScheme/a:minorFont", ns)
            if minor is not None:
                latin = minor.find("a:latin", ns)
                ea = minor.find("a:ea", ns)
                if latin is not None:
                    self._theme_fonts.font_ascii = latin.get("typeface")
                    self._theme_fonts.font_hAnsi = latin.get("typeface")
                if ea is not None and ea.get("typeface"):
                    self._theme_fonts.font_east_asia = ea.get("typeface")
                if not self._theme_fonts.font_east_asia:
                    for sup in minor.findall("a:font", ns):
                        if sup.get("script", "") in ("Hans", "Hant"):
                            self._theme_fonts.font_east_asia = sup.get("typeface")
                            break

            # ── Issue #4: 主题颜色 ──
            clr_scheme = root.find(".//a:themeElements/a:clrScheme", ns)
            if clr_scheme is not None:
                for color_elem in clr_scheme:
                    name = etree.QName(color_elem.tag).localname
                    # 子元素可能是 a:srgbClr 或 a:sysClr
                    srgb = color_elem.find("a:srgbClr", ns)
                    if srgb is not None:
                        val = srgb.get("val", "")
                        if val:
                            self._theme_colors[name] = val.upper()
                    else:
                        sys_clr = color_elem.find("a:sysClr", ns)
                        if sys_clr is not None:
                            last_clr = sys_clr.get("lastClr", "")
                            if last_clr:
                                self._theme_colors[name] = last_clr.upper()
        except Exception:
            pass

    def _parse_doc_defaults(self) -> None:
        try:
            styles_elem = self.document.part._styles_part._element
            doc_defaults = styles_elem.find(qn("w:docDefaults"))
            if doc_defaults is None:
                return
            rpr_def = doc_defaults.find(qn("w:rPrDefault"))
            if rpr_def is not None:
                rpr = rpr_def.find(qn("w:rPr"))
                if rpr is not None:
                    self._doc_defaults_rpr = self._parse_rpr(rpr)
            ppr_def = doc_defaults.find(qn("w:pPrDefault"))
            if ppr_def is not None:
                ppr = ppr_def.find(qn("w:pPr"))
                if ppr is not None:
                    self._doc_defaults_ppr = self._parse_ppr(ppr)
        except Exception:
            pass

    # ── Issue #2: 解析 numbering.xml ──

    def _parse_numbering(self) -> None:
        """解析 numbering.xml，建立 numId -> ilvl -> {run_style, para_style} 映射"""
        try:
            with ZipFile(str(self.file_path), "r") as zf:
                if "word/numbering.xml" not in zf.namelist():
                    return
                root = etree.fromstring(zf.read("word/numbering.xml"))
        except Exception:
            return

        # 先建立 abstractNumId -> abstractNum 元素映射
        abstract_map: dict[str, etree._Element] = {}
        for an in root.findall(qn("w:abstractNum")):
            aid = an.get(qn("w:abstractNumId"), "")
            if aid:
                abstract_map[aid] = an

        # 建立 numId -> abstractNumId 映射，并收集 lvlOverride
        for num_elem in root.findall(qn("w:num")):
            nid_str = num_elem.get(qn("w:numId"), "")
            if not nid_str.isdigit():
                continue
            nid = int(nid_str)
            an_ref = num_elem.find(qn("w:abstractNumId"))
            if an_ref is None:
                continue
            aid = an_ref.get(qn("w:val"), "")
            an_elem = abstract_map.get(aid)
            if an_elem is None:
                continue

            lvl_styles: dict[int, dict] = {}

            # 从 abstractNum 读取每个 lvl
            for lvl in an_elem.findall(qn("w:lvl")):
                ilvl_str = lvl.get(qn("w:ilvl"), "0")
                ilvl = int(ilvl_str) if ilvl_str.isdigit() else 0
                rs = RunStyle()
                ps = ParagraphStyle()
                rpr = lvl.find(qn("w:rPr"))
                if rpr is not None:
                    rs = self._parse_rpr(rpr)
                ppr = lvl.find(qn("w:pPr"))
                if ppr is not None:
                    ps = self._parse_ppr(ppr)
                lvl_styles[ilvl] = {"run": rs, "para": ps}

            # lvlOverride 可以覆盖特定 ilvl
            for ovr in num_elem.findall(qn("w:lvlOverride")):
                ilvl_str = ovr.get(qn("w:ilvl"), "0")
                ilvl = int(ilvl_str) if ilvl_str.isdigit() else 0
                lvl_elem = ovr.find(qn("w:lvl"))
                if lvl_elem is not None:
                    rs = RunStyle()
                    ps = ParagraphStyle()
                    rpr = lvl_elem.find(qn("w:rPr"))
                    if rpr is not None:
                        rs = self._parse_rpr(rpr)
                    ppr = lvl_elem.find(qn("w:pPr"))
                    if ppr is not None:
                        ps = self._parse_ppr(ppr)
                    lvl_styles[ilvl] = {"run": rs, "para": ps}

            self._numbering_cache[nid] = lvl_styles

    # ── rPr 解析（精确四槽位） ──

    def _parse_rpr(self, rpr: etree._Element) -> RunStyle:
        s = RunStyle()

        # ── 字体四槽位 + Issue #5: 完整 fallback ──
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            # 直接值优先
            s.font_ascii = rfonts.get(qn("w:ascii"))
            s.font_east_asia = rfonts.get(qn("w:eastAsia"))
            s.font_hAnsi = rfonts.get(qn("w:hAnsi"))
            s.font_cs = rfonts.get(qn("w:cs"))

            # theme 回退（asciiTheme / hAnsiTheme / eastAsiaTheme / cstheme）
            if not s.font_ascii:
                theme_key = rfonts.get(qn("w:asciiTheme"), "")
                if theme_key:
                    s.font_ascii = self._resolve_theme_font(theme_key)
            if not s.font_hAnsi:
                theme_key = rfonts.get(qn("w:hAnsiTheme"), "")
                if theme_key:
                    s.font_hAnsi = self._resolve_theme_font(theme_key)
            if not s.font_east_asia:
                theme_key = rfonts.get(qn("w:eastAsiaTheme"), "")
                if theme_key:
                    s.font_east_asia = self._resolve_theme_font(theme_key)
            if not s.font_cs:
                theme_key = rfonts.get(qn("w:cstheme"), "")
                if theme_key:
                    s.font_cs = self._resolve_theme_font(theme_key)

            # Issue #5: hAnsi fallback → ascii
            if not s.font_hAnsi and s.font_ascii:
                s.font_hAnsi = s.font_ascii

        sz = rpr.find(qn("w:sz"))
        s.font_size = _half_pt(sz.get(qn("w:val")) if sz is not None else None)
        sz_cs = rpr.find(qn("w:szCs"))
        s.font_size_cs = _half_pt(sz_cs.get(qn("w:val")) if sz_cs is not None else None)

        # ── Issue #1: Toggle 字段记录原始值用于翻转继承 ──
        _toggle_map = {
            "bold":       qn("w:b"),
            "bold_cs":    qn("w:bCs"),
            "italic":     qn("w:i"),
            "italic_cs":  qn("w:iCs"),
            "strike":     qn("w:strike"),
            "dstrike":    qn("w:dstrike"),
            "outline":    qn("w:outline"),
            "shadow":     qn("w:shadow"),
            "emboss":     qn("w:emboss"),
            "imprint":    qn("w:imprint"),
            "vanish":     qn("w:vanish"),
            "small_caps": qn("w:smallCaps"),
            "all_caps":   qn("w:caps"),
        }
        for field_name, tag in _toggle_map.items():
            elem = rpr.find(tag)
            raw_val = _get_toggle_raw(elem)
            setattr(s, field_name, raw_val)
            if raw_val is not None:
                s._toggle_raw[field_name] = raw_val

        u = rpr.find(qn("w:u"))
        if u is not None:
            s.underline = u.get(qn("w:val"), "single")

        # ── Issue #4: 颜色 + themeColor 解析 ──
        color = rpr.find(qn("w:color"))
        if color is not None:
            val = color.get(qn("w:val"), "")
            theme_color = color.get(qn("w:themeColor"), "")
            theme_tint = color.get(qn("w:themeTint"), "")
            theme_shade = color.get(qn("w:themeShade"), "")
            if val and val.lower() != "auto":
                s.color = f"#{val.upper()}"
            elif theme_color:
                resolved = self._resolve_theme_color(theme_color, theme_tint, theme_shade)
                if resolved:
                    s.color = resolved

        hl = rpr.find(qn("w:highlight"))
        if hl is not None:
            s.highlight = hl.get(qn("w:val"))

        shd = rpr.find(qn("w:shd"))
        if shd is not None:
            fill = shd.get(qn("w:fill"), "")
            if fill and fill.lower() not in ("auto", ""):
                s.shading = f"#{fill.upper()}"

        sp = rpr.find(qn("w:spacing"))
        s.spacing = _twips(sp.get(qn("w:val")) if sp is not None else None)

        kern = rpr.find(qn("w:kern"))
        s.kern = _half_pt(kern.get(qn("w:val")) if kern is not None else None)

        pos = rpr.find(qn("w:position"))
        s.position = _half_pt(pos.get(qn("w:val")) if pos is not None else None)

        # ── Issue #8: w:w (字符缩放) 和 w:fitText ──
        w_scale = rpr.find(qn("w:w"))
        if w_scale is not None:
            v = w_scale.get(qn("w:val"), "")
            if v.isdigit():
                s.char_scale = int(v)

        fit_text = rpr.find(qn("w:fitText"))
        if fit_text is not None:
            v = fit_text.get(qn("w:val"), "")
            if v.isdigit():
                s.fit_text = int(v)

        va = rpr.find(qn("w:vertAlign"))
        if va is not None:
            s.vertical_align = va.get(qn("w:val"))

        lang = rpr.find(qn("w:lang"))
        if lang is not None:
            s.lang_val = lang.get(qn("w:val"))
            s.lang_east_asia = lang.get(qn("w:eastAsia"))
            s.lang_bidi = lang.get(qn("w:bidi"))

        return s

    def _resolve_theme_font(self, theme_key: str) -> Optional[str]:
        """将 asciiTheme/hAnsiTheme 等 key 映射到实际字体名"""
        # theme_key 形如 "minorHAnsi", "majorEastAsia" 等
        key_lower = theme_key.lower()
        if "eastasia" in key_lower or "eastAsia" in theme_key:
            return self._theme_fonts.font_east_asia
        if "cs" in key_lower:
            return self._theme_fonts.font_cs
        # minor/major → ascii/hAnsi
        return self._theme_fonts.font_ascii

    def _resolve_theme_color(self, theme_color: str, tint: str = "", shade: str = "") -> Optional[str]:
        """Issue #4: 将 themeColor 名称解析为 #RRGGBB 字符串"""
        # theme_color 形如 "accent1", "dk1", "lt1" 等
        # styles.xml 中的名称与 theme1.xml clrScheme 子元素名称对应
        # 常见别名映射
        alias = {
            "dk1": "dk1", "lt1": "lt1", "dk2": "dk2", "lt2": "lt2",
            "accent1": "accent1", "accent2": "accent2", "accent3": "accent3",
            "accent4": "accent4", "accent5": "accent5", "accent6": "accent6",
            "hlink": "hlink", "folHlink": "folHlink",
        }
        key = alias.get(theme_color, theme_color)
        hex_val = self._theme_colors.get(key)
        if not hex_val:
            return None

        # 应用 tint / shade（简化计算：线性混合到白/黑）
        try:
            r = int(hex_val[0:2], 16)
            g = int(hex_val[2:4], 16)
            b = int(hex_val[4:6], 16)
            if tint:
                t = int(tint, 16) / 255.0
                r = int(r + (255 - r) * (1 - t))
                g = int(g + (255 - g) * (1 - t))
                b = int(b + (255 - b) * (1 - t))
            elif shade:
                s_val = int(shade, 16) / 255.0
                r = int(r * s_val)
                g = int(g * s_val)
                b = int(b * s_val)
            return f"#{r:02X}{g:02X}{b:02X}"
        except Exception:
            return f"#{hex_val}"

    # ── pPr 解析 ──

    def _parse_ppr(self, ppr: etree._Element) -> ParagraphStyle:
        s = ParagraphStyle()
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            s.alignment = jc.get(qn("w:val"))

        olvl = ppr.find(qn("w:outlineLvl"))
        if olvl is not None:
            v = olvl.get(qn("w:val"), "")
            if v.isdigit():
                s.outline_level = int(v)

        sp = ppr.find(qn("w:spacing"))
        if sp is not None:
            line = sp.get(qn("w:line"))
            rule = sp.get(qn("w:lineRule"), "auto")
            s.line_spacing_rule = rule
            if line:
                try:
                    lv = int(line)
                    s.line_spacing = lv / 240.0 if rule == "auto" else lv / 20.0
                except ValueError:
                    pass
            s.space_before = _twips(sp.get(qn("w:before")))
            s.space_after = _twips(sp.get(qn("w:after")))

        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            s.indent_left = _twips(ind.get(qn("w:left")) or ind.get(qn("w:start")))
            s.indent_right = _twips(ind.get(qn("w:right")) or ind.get(qn("w:end")))
            s.indent_first_line = _twips(ind.get(qn("w:firstLine")))
            s.indent_hanging = _twips(ind.get(qn("w:hanging")))

        s.keep_next = _get_toggle(ppr.find(qn("w:keepNext")))
        s.keep_lines = _get_toggle(ppr.find(qn("w:keepLines")))
        s.page_break_before = _get_toggle(ppr.find(qn("w:pageBreakBefore")))

        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            numid_elem = numpr.find(qn("w:numId"))
            ilvl_elem = numpr.find(qn("w:ilvl"))
            if numid_elem is not None:
                nv = numid_elem.get(qn("w:val"), "0")
                if nv.lstrip("-").isdigit() and int(nv) != 0:
                    s.num_id = int(nv)
                    s.num_id_is_direct = True
            if ilvl_elem is not None:
                iv = ilvl_elem.get(qn("w:val"), "0")
                if iv.lstrip("-").isdigit():
                    s.num_ilvl = int(iv)

        inner_rpr = ppr.find(qn("w:rPr"))
        if inner_rpr is not None:
            s.embedded_rpr = self._parse_rpr(inner_rpr)

        # 捕获分节符 (Section Break)
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is not None:
            s.section_xml = _elem_to_xml_str(sect_pr)
            self._parse_sect_pr_layout(sect_pr, s)

        return s

    def _parse_sect_pr_layout(self, sect_pr: etree._Element, ps: ParagraphStyle) -> None:
        """Issue #6: 从 sectPr 提取页面尺寸、页边距、分栏数"""
        pg_sz = sect_pr.find(qn("w:pgSz"))
        if pg_sz is not None:
            w = pg_sz.get(qn("w:w"), "")
            h = pg_sz.get(qn("w:h"), "")
            if w.isdigit():
                ps.page_width = int(w) / 20.0   # twips -> 磅
            if h.isdigit():
                ps.page_height = int(h) / 20.0

        pg_mar = sect_pr.find(qn("w:pgMar"))
        if pg_mar is not None:
            left = pg_mar.get(qn("w:left"), "")
            right = pg_mar.get(qn("w:right"), "")
            if left.isdigit():
                ps.page_margin_left = int(left) / 20.0
            if right.isdigit():
                ps.page_margin_right = int(right) / 20.0

        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"), "1")
            if num.isdigit():
                ps.page_cols = int(num)

    # ── 样式继承链 ──

    def _resolve_rpr_chain(self, style_id: str, visited: Optional[set] = None) -> RunStyle:
        if style_id in self._resolved_run_cache:
            return copy.deepcopy(self._resolved_run_cache[style_id])
        if visited is None:
            visited = set()
        if style_id in visited:
            return RunStyle()
        visited.add(style_id)

        elem = self._style_cache.get(style_id)
        if elem is None:
            return RunStyle()

        cur = RunStyle()
        rpr = elem.find(qn("w:rPr"))
        if rpr is not None:
            cur = self._parse_rpr(rpr)
        ppr = elem.find(qn("w:pPr"))
        if ppr is not None:
            inner = ppr.find(qn("w:rPr"))
            if inner is not None:
                cur.merge_from(self._parse_rpr(inner))

        based_on = elem.find(qn("w:basedOn"))
        if based_on is not None:
            pid = based_on.get(qn("w:val"), "")
            if pid:
                cur.merge_from(self._resolve_rpr_chain(pid, visited))

        self._resolved_run_cache[style_id] = copy.deepcopy(cur)
        return cur

    def _resolve_ppr_chain(self, style_id: str, visited: Optional[set] = None) -> ParagraphStyle:
        if style_id in self._resolved_para_cache:
            return copy.deepcopy(self._resolved_para_cache[style_id])
        if visited is None:
            visited = set()
        if style_id in visited:
            return ParagraphStyle()
        visited.add(style_id)

        elem = self._style_cache.get(style_id)
        if elem is None:
            return ParagraphStyle()

        cur = ParagraphStyle()
        ppr = elem.find(qn("w:pPr"))
        if ppr is not None:
            cur = self._parse_ppr(ppr)

        based_on = elem.find(qn("w:basedOn"))
        if based_on is not None:
            pid = based_on.get(qn("w:val"), "")
            if pid:
                cur.merge_from(self._resolve_ppr_chain(pid, visited))

        self._resolved_para_cache[style_id] = copy.deepcopy(cur)
        return cur

    # ── 计算最终样式 ──

    def _compute_run_style(
        self, run_elem: etree._Element, para_style_id: Optional[str],
        para_num_id: Optional[int] = None, para_num_ilvl: int = 0,
        table_style_id: Optional[str] = None,
    ) -> tuple[RunStyle, dict[str, dict], RunStyle]:
        layers: dict[str, dict] = {}

        # 第5层 theme
        l5 = copy.deepcopy(self._theme_fonts)
        layers["5_theme"] = l5.to_dict()

        # 第4层 docDefaults
        l4 = copy.deepcopy(self._doc_defaults_rpr)
        l4.merge_from(l5)
        layers["4_docDefaults"] = self._doc_defaults_rpr.to_dict()

        # Issue #2: numbering run 样式（优先级低于段落样式，高于 docDefaults）
        l_num = RunStyle()
        if para_num_id and para_num_id in self._numbering_cache:
            lvl_data = self._numbering_cache[para_num_id].get(para_num_ilvl, {})
            l_num = copy.deepcopy(lvl_data.get("run", RunStyle()))
        l_num.merge_from(l4)
        layers["4b_numbering"] = l_num.to_dict() if para_num_id else {}

        # Issue #3: 表格样式（优先级低于段落样式，高于 numbering）
        l_tbl = RunStyle()
        if table_style_id:
            l_tbl = self._resolve_rpr_chain(table_style_id)
        l_tbl.merge_from(l_num)
        layers["4c_tableStyle"] = l_tbl.to_dict() if table_style_id else {}

        # 第3层 段落样式链（无样式时回退到 Normal）
        effective_para_sid = para_style_id or self._normal_style_id
        l3 = RunStyle()
        if effective_para_sid:
            l3 = self._resolve_rpr_chain(effective_para_sid)
        l3.merge_from(l_tbl)
        layers["3_paragraphStyle"] = (
            self._resolve_rpr_chain(effective_para_sid).to_dict() if effective_para_sid else {}
        )

        # 第2层 字符样式
        l2 = RunStyle()
        char_sid: Optional[str] = None
        rpr = run_elem.find(qn("w:rPr"))
        if rpr is not None:
            rs = rpr.find(qn("w:rStyle"))
            if rs is not None:
                char_sid = rs.get(qn("w:val"))
                if char_sid:
                    l2 = self._resolve_rpr_chain(char_sid)
        l2.merge_from(l3)
        layers["2_characterStyle"] = (
            self._resolve_rpr_chain(char_sid).to_dict() if char_sid else {}
        )

        # 第1层 直接格式
        l1 = RunStyle()
        if rpr is not None:
            l1 = self._parse_rpr(rpr)
        l1.merge_from(l2)
        layers["1_directFormat"] = self._parse_rpr(rpr).to_dict() if rpr is not None else {}

        direct_only = self._parse_rpr(rpr) if rpr is not None else RunStyle()
        return l1, layers, direct_only

    def _compute_para_style(
        self, para_elem: etree._Element
    ) -> tuple[ParagraphStyle, Optional[str]]:
        ppr = para_elem.find(qn("w:pPr"))
        psid: Optional[str] = None
        direct = ParagraphStyle()
        if ppr is not None:
            ps = ppr.find(qn("w:pStyle"))
            if ps is not None:
                psid = ps.get(qn("w:val"))
            direct = self._parse_ppr(ppr)

        effective_psid = psid or self._normal_style_id
        chain = ParagraphStyle()
        if effective_psid:
            chain = self._resolve_ppr_chain(effective_psid)
        chain.merge_from(copy.deepcopy(self._doc_defaults_ppr))
        direct.merge_from(chain)

        # Issue #2: 合并 numbering 样式（列表缩进/字体）
        num_id = direct.num_id
        num_ilvl = direct.num_ilvl if direct.num_ilvl is not None else 0
        if num_id and num_id in self._numbering_cache:
            lvl_data = self._numbering_cache[num_id].get(num_ilvl, {})
            num_para: ParagraphStyle = lvl_data.get("para", ParagraphStyle())
            # numbering 样式优先级低于直接格式，仅填充未设置的字段
            direct.merge_from(num_para)

        return direct, psid

    def _get_style_name(self, sid: str) -> Optional[str]:
        if sid in self._style_cache:
            ne = self._style_cache[sid].find(qn("w:name"))
            if ne is not None:
                return ne.get(qn("w:val"))
        return sid

    # ── 嵌入对象检测 ──

    def _run_has_embedded(self, run_elem: etree._Element) -> Optional[EmbeddedObject]:
        """检测 Run 中是否包含 Drawing/图片/图表/OLE/公式等非文本对象"""
        # 定义命名空间
        ns = {
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006"
        }

        # 检查 w:r 内部的 m:oMath（行内公式包裹在 run 里的情况）
        M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        omath_in_run = run_elem.find(f"{{{M_NS}}}oMath")
        if omath_in_run is not None:
            return EmbeddedObject(
                object_type="math",
                xml_snippet=_elem_to_xml_str(run_elem),
                relationship_ids=[],
            )

        mc_ac = run_elem.find(".//mc:AlternateContent", namespaces=ns)
        # 检查 drawing（图片/图表/SmartArt）
        drawing = run_elem.find(qn("w:drawing"))
        if drawing is not None:
            obj_type = "drawing"
            # 进一步判断是否为 chart
            chart_refs = drawing.findall(".//" + qn("c:chart"))
            if not chart_refs:
                # 尝试用通配符搜索
                for desc in drawing.iter():
                    tag = etree.QName(desc.tag).localname if isinstance(desc.tag, str) else ""
                    if tag == "chart":
                        chart_refs = [desc]
                        break
            if chart_refs:
                obj_type = "chart"

            rids = _find_relationship_ids(drawing)
            return EmbeddedObject(
                object_type=obj_type,
                xml_snippet=_elem_to_xml_str(run_elem),
                relationship_ids=rids,
            )

        # 检查 w:object（OLE 嵌入）
        obj = run_elem.find(qn("w:object"))
        if obj is not None:
            rids = _find_relationship_ids(obj)
            return EmbeddedObject(
                object_type="ole",
                xml_snippet=_elem_to_xml_str(run_elem),
                relationship_ids=rids,
            )

        # 检查 w:pict（VML 图片/旧式图形）
        pict = run_elem.find(qn("w:pict"))
        if pict is not None:
            rids = _find_relationship_ids(pict)
            return EmbeddedObject(
                object_type="pict",
                xml_snippet=_elem_to_xml_str(run_elem),
                relationship_ids=rids,
            )

        # 检查 mc:AlternateContent（兼容性包装）
        mc_ac = run_elem.find(".//mc:AlternateContent", namespaces=ALL_NS)
        if mc_ac is not None:
            rids = _find_relationship_ids(mc_ac)
            return EmbeddedObject(
                object_type="alternateContent",
                xml_snippet=_elem_to_xml_str(run_elem),
                relationship_ids=rids,
            )

        return None

    def _para_has_embedded_only(self, para_elem: etree._Element) -> Optional[EmbeddedObject]:
        """检查段落级别的嵌入对象（如独立的 mc:AlternateContent）"""
        for child in para_elem:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
            if tag in ("bookmarkStart", "bookmarkEnd", "proofErr", "pPr"):
                continue
            if tag == "r":
                continue  # Run 在别处处理
            # 段落直接子级的非标准元素
            if tag in ("ins", "del", "moveFrom", "moveTo"):
                continue
        return None

    # ── 核心提取 ──

    def extract_all(self) -> list[TextSegment]:
        """提取所有文本片段和嵌入对象"""
        segments: list[TextSegment] = []
        idx = 0
        body = self.document.element.body

        for child in body:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
            match tag:
                case "p":
                    idx = self._process_para(child, "正文", segments, idx, None)
                case "tbl":
                    idx = self._process_table(child, segments, idx)
                case "sdt":
                    idx = self._process_sdt(child, "正文", segments, idx, None)
                case "sectPr":
                    # 捕获文档末尾的分节符 (通常是最后一节的布局配置)
                    if segments:
                        last_seg = segments[-1]
                        if last_seg.computed_para_style:
                            last_seg.computed_para_style.section_xml = _elem_to_xml_str(child)
                            # Issue #6: 解析页面布局参数
                            self._parse_sect_pr_layout(child, last_seg.computed_para_style)

        # 页眉
        for sec in self.document.sections:
            try:
                hdr = sec.header
                if hdr and not hdr.is_linked_to_previous:
                    for p in hdr.paragraphs:
                        idx = self._process_para(p._element, "页眉", segments, idx, None)
            except Exception:
                pass

        # 页脚
        for sec in self.document.sections:
            try:
                ftr = sec.footer
                if ftr and not ftr.is_linked_to_previous:
                    for p in ftr.paragraphs:
                        idx = self._process_para(p._element, "页脚", segments, idx, None)
            except Exception:
                pass

        return segments

    def _process_para(
        self,
        para_elem: etree._Element,
        source: str,
        segments: list[TextSegment],
        idx: int,
        table_meta: Optional[TableMeta],
    ) -> int:
        computed_para, psid = self._compute_para_style(para_elem)
        pname = self._get_style_name(psid) if psid else None
        self._para_id_counter += 1
        pid = self._para_id_counter
        raw_para_xml = _elem_to_xml_str(para_elem)
        first_seg_in_para: Optional[int] = None

        has_content = False

        M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

        for run_elem in para_elem:
            child_tag = etree.QName(run_elem.tag).localname if isinstance(run_elem.tag, str) else ""
            child_ns = etree.QName(run_elem.tag).namespace if isinstance(run_elem.tag, str) else ""
            if child_tag == "pPr":
                continue

            # 公式：m:oMath 或 m:oMathPara 是段落直接子元素
            if child_ns == M_NS and child_tag in ("oMath", "oMathPara"):
                idx += 1
                has_content = True
                segments.append(TextSegment(
                    index=idx,
                    source=source,
                    text="",
                    paragraph_style_name=pname,
                    computed_para_style=computed_para,
                    computed_run_style=RunStyle(),
                    table_meta=copy.deepcopy(table_meta) if table_meta else None,
                    paragraph_id=pid,
                    is_embedded_object=True,
                    embedded_object=EmbeddedObject(
                        object_type="math",
                        xml_snippet=_elem_to_xml_str(run_elem),
                        relationship_ids=[],
                    ),
                ))
                if first_seg_in_para is None:
                    first_seg_in_para = len(segments) - 1
                continue

            if child_tag != "r":
                continue

            # 检查是否包含域字符 (fldChar/instrText) —— 页码等域
            has_fld = (
                run_elem.find(qn("w:fldChar")) is not None
                or run_elem.find(qn("w:instrText")) is not None
            )
            if has_fld:
                idx += 1
                has_content = True
                rids = _find_relationship_ids(run_elem)
                segments.append(TextSegment(
                    index=idx,
                    source=source,
                    text="",
                    paragraph_style_name=pname,
                    computed_para_style=computed_para,
                    computed_run_style=RunStyle(),
                    table_meta=copy.deepcopy(table_meta) if table_meta else None,
                    paragraph_id=pid,
                    is_embedded_object=True,
                    embedded_object=EmbeddedObject(
                        object_type="fldChar",
                        xml_snippet=_elem_to_xml_str(run_elem),
                        relationship_ids=rids,
                    ),
                ))
                if first_seg_in_para is None:
                    first_seg_in_para = len(segments) - 1
                continue

            # 检查嵌入对象
            emb = self._run_has_embedded(run_elem)
            if emb is not None:
                idx += 1
                has_content = True
                segments.append(TextSegment(
                    index=idx,
                    source=source,
                    text="",
                    paragraph_style_name=pname,
                    computed_para_style=computed_para,
                    computed_run_style=RunStyle(),
                    table_meta=copy.deepcopy(table_meta) if table_meta else None,
                    paragraph_id=pid,
                    is_embedded_object=True,
                    embedded_object=emb,
                ))
                if first_seg_in_para is None:
                    first_seg_in_para = len(segments) - 1
                continue

            # 普通文本 Run
            texts = []
            for t in run_elem.iterchildren(qn("w:t")):
                if t.text:
                    texts.append(t.text)
            text = "".join(texts)

            if not text:
                continue

            has_content = True
            computed_run, layers, direct_run = self._compute_run_style(
                run_elem, psid,
                para_num_id=computed_para.num_id,
                para_num_ilvl=computed_para.num_ilvl or 0,
                table_style_id=table_meta.table_style_name if table_meta else None,
            )

            # 字符样式名
            cname: Optional[str] = None
            rpr = run_elem.find(qn("w:rPr"))
            if rpr is not None:
                rs = rpr.find(qn("w:rStyle"))
                if rs is not None:
                    cid = rs.get(qn("w:val"), "")
                    if cid:
                        cname = self._get_style_name(cid)

            idx += 1
            segments.append(TextSegment(
                index=idx,
                source=source,
                text=text,
                paragraph_style_name=pname,
                character_style_name=cname,
                computed_run_style=computed_run,
                direct_run_style=direct_run,
                computed_para_style=computed_para,
                style_layers=layers,
                table_meta=copy.deepcopy(table_meta) if table_meta else None,
                paragraph_id=pid,
            ))
            if first_seg_in_para is None:
                first_seg_in_para = len(segments) - 1

        # 空段落保留结构
        if not has_content:
            idx += 1
            segments.append(TextSegment(
                index=idx,
                source=source,
                text="",
                paragraph_style_name=pname,
                computed_para_style=computed_para,
                computed_run_style=RunStyle(),
                table_meta=copy.deepcopy(table_meta) if table_meta else None,
                paragraph_id=pid,
                raw_paragraph_xml=raw_para_xml,
            ))
        elif first_seg_in_para is not None:
            segments[first_seg_in_para].raw_paragraph_xml = raw_para_xml

        return idx

    def _process_table(
        self, tbl_elem: etree._Element, segments: list[TextSegment], idx: int
    ) -> int:
        self._table_id_counter += 1
        tid = self._table_id_counter

        table_xml = _elem_to_xml_str(tbl_elem)

        # 提取表格整体样式 (tblPr)
        tbl_pr_xml = None
        tbl_style_name: Optional[str] = None
        tbl_pr = tbl_elem.find(qn("w:tblPr"))
        if tbl_pr is not None:
            tbl_pr_xml = _elem_to_xml_str(tbl_pr)
            # Issue #3: 提取 tblStyle 名称
            tbl_style_elem = tbl_pr.find(qn("w:tblStyle"))
            if tbl_style_elem is not None:
                tbl_style_name = tbl_style_elem.get(qn("w:val"))

        rows = list(tbl_elem.iterchildren(qn("w:tr")))
        total_rows = len(rows)
        max_cols = 0
        for tr in rows:
            max_cols = max(max_cols, len(list(tr.iterchildren(qn("w:tc")))))

        for ri, tr in enumerate(rows):
            cells = list(tr.iterchildren(qn("w:tc")))
            for ci, tc in enumerate(cells):
                meta = TableMeta(
                    table_id=tid, row=ri, col=ci,
                    total_rows=total_rows, total_cols=max_cols,
                    table_pr_xml=tbl_pr_xml,
                    table_xml=table_xml,
                    table_style_name=tbl_style_name,
                )
                
                tcpr = tc.find(qn("w:tcPr"))
                if tcpr is not None:
                    # 存储单元格样式 (tcPr)
                    meta.cell_pr_xml = _elem_to_xml_str(tcpr)
                    
                    # 提取宽度
                    tw = tcpr.find(qn("w:tcW"))
                    if tw is not None:
                        wv = tw.get(qn("w:w"), "")
                        wt = tw.get(qn("w:type"), "dxa")
                        if wv.isdigit() and wt == "dxa":
                            meta.cell_width = int(wv) / 20.0
                    
                    # 提取合并信息
                    gs = tcpr.find(qn("w:gridSpan"))
                    if gs is not None:
                        gv = gs.get(qn("w:val"), "1")
                        if gv.isdigit():
                            meta.grid_span = int(gv)
                    
                    vm = tcpr.find(qn("w:vMerge"))
                    if vm is not None:
                        meta.v_merge = vm.get(qn("w:val"), "continue")

                for child in tc:
                    ctag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
                    match ctag:
                        case "p":
                            idx = self._process_para(child, "表格", segments, idx, meta)
                        case "tbl":
                            idx = self._process_table(child, segments, idx)
                        case "sdt":
                            idx = self._process_sdt(child, "表格", segments, idx, meta)
        return idx

    def _process_sdt(
        self,
        sdt_elem: etree._Element,
        source: str,
        segments: list[TextSegment],
        idx: int,
        table_meta: Optional[TableMeta],
    ) -> int:
        self._para_id_counter += 1
        pid = self._para_id_counter
        idx += 1
        rids = _find_relationship_ids(sdt_elem)
        segments.append(TextSegment(
            index=idx,
            source=source,
            text="",
            computed_para_style=ParagraphStyle(),
            computed_run_style=RunStyle(),
            table_meta=copy.deepcopy(table_meta) if table_meta else None,
            paragraph_id=pid,
            is_embedded_object=True,
            embedded_object=EmbeddedObject(
                object_type="sdt",
                xml_snippet=_elem_to_xml_str(sdt_elem),
                relationship_ids=rids,
            ),
        ))
        return idx

if __name__ == "__main__":

    def to_dict_recursive(obj):
        """递归地将对象及其嵌套属性转换为字典/列表/基础类型"""
        if isinstance(obj, list):
            return [to_dict_recursive(item) for item in obj]
        if isinstance(obj, dict):
            return {k: to_dict_recursive(v) for k, v in obj.items()}
        if hasattr(obj, "__dict__"):
            return to_dict_recursive(obj.__dict__)
        # 如果是基础类型 (str, int, float, bool, None)，直接返回
        return obj
    target = r"C:\Users\H\Desktop\数检_程序-AI\测试文件\原文-含不可编辑_01 (2026-007)2025年年度报告.docx"
    e=EnhancedExtractorV2(target)
    result = e.extract_all()
    # 使用递归函数处理整个嵌套结构
    result_serializable = to_dict_recursive(result)

    # 现在序列化就不会出错了
    formatted = json.dumps(result_serializable, indent=4, ensure_ascii=False)
    print(formatted)