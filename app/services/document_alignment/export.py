from __future__ import annotations

import csv
from copy import copy
from dataclasses import dataclass
from difflib import SequenceMatcher
from io import BytesIO, StringIO
import json
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import BadZipFile, ZipFile

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy.orm import Session

from app.models import (
    DocumentAlignmentPair, FileRecord, ProofreadingBatch, ProofreadingColumnBinding,
    ProofreadingSegmentBaseline, Segment,
)
from app.services.normalizer import compact_match_core

from .parser import AlignUnit, parse_side
from .segments import TRANSLATION_ONLY_SOURCE_LABEL, ensure_document_pair_segments_complete
from .service import target_cache_path

MISSING_TRANSLATION_LABEL = "【译文缺失】"
TRANSLATION_ONLY_EXPORT_LABEL = "【增译】"


@dataclass(frozen=True)
class ProofreadingExportRow:
    order: int
    kind: str
    source_text: str
    original_target_text: str
    reviewed_target_text: str
    changed: bool
    confirmation_status: str
    llm_status: str
    confidence: float | None
    method: str
    pair_id: str = ""
    block_type: str = "paragraph"
    block_index: int = 0
    row_index: int | None = None
    cell_index: int | None = None


def _segment_metadata(segment: Segment) -> dict[str, Any]:
    try:
        value = json.loads(segment.segment_metadata or "{}")
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def build_proofreading_export_rows(
    db: Session, batch: ProofreadingBatch,
) -> tuple[list[ProofreadingExportRow], FileRecord | None]:
    """建立导出的唯一顺序清单，不依赖 sentence_id 或数据库自然顺序。"""
    if batch.batch_kind == "document_pair":
        ensure_document_pair_segments_complete(db, batch)
        binding = db.query(ProofreadingColumnBinding).filter_by(batch_id=batch.id).first()
        file_record = db.get(FileRecord, binding.file_record_id) if binding else None
        baselines = db.query(ProofreadingSegmentBaseline).filter_by(batch_id=batch.id).all()
        baseline_by_segment = {item.segment_id: item for item in baselines}
        segment_by_pair: dict[str, Segment] = {}
        if file_record is not None:
            for segment in db.query(Segment).filter_by(file_record_id=file_record.id).all():
                pair_id = str(_segment_metadata(segment).get("alignment_pair_id") or "")
                if pair_id:
                    segment_by_pair[pair_id] = segment
        rows: list[ProofreadingExportRow] = []
        pairs = db.query(DocumentAlignmentPair).filter_by(batch_id=batch.id).order_by(
            DocumentAlignmentPair.pair_order,
        ).all()
        for pair in pairs:
            segment = segment_by_pair.get(str(pair.id))
            baseline = baseline_by_segment.get(segment.id) if segment is not None else None
            source_text = segment.source_text if segment is not None else pair.source_text
            reviewed_target = segment.target_text if segment is not None else pair.target_text
            original_target = baseline.original_target_text if baseline is not None else pair.target_text
            try:
                source_indices = json.loads(pair.src_indices or "[]")
            except (TypeError, ValueError, json.JSONDecodeError):
                source_indices = []
            translation_only = not bool(source_indices if isinstance(source_indices, list) else [])
            kind = "增译" if translation_only else "缺译" if not reviewed_target else "对齐"
            llm_status = (
                "已校对" if segment is not None and bool(segment.llm_provider)
                else "人工确认" if segment is not None and segment.status == "confirmed"
                else "未校对"
            )
            rows.append(ProofreadingExportRow(
                order=pair.pair_order,
                kind=kind,
                source_text=TRANSLATION_ONLY_SOURCE_LABEL if translation_only else source_text,
                original_target_text=original_target or "",
                reviewed_target_text=reviewed_target or "",
                changed=(reviewed_target or "") != (original_target or ""),
                confirmation_status="已确认" if segment is not None and segment.status == "confirmed" else "未确认",
                llm_status=llm_status,
                confidence=pair.confidence,
                method=pair.method,
                pair_id=str(pair.id),
                block_type=pair.block_type,
                block_index=pair.block_index,
                row_index=pair.row_index,
                cell_index=pair.cell_index,
            ))
        return rows, file_record

    bindings = db.query(ProofreadingColumnBinding).filter_by(batch_id=batch.id).all()
    file_record = db.get(FileRecord, bindings[0].file_record_id) if bindings else None
    ordered_items: list[tuple[int, int, int, int, ProofreadingSegmentBaseline, Segment]] = []
    for binding in bindings:
        items = db.query(ProofreadingSegmentBaseline, Segment).join(
            Segment, Segment.id == ProofreadingSegmentBaseline.segment_id,
        ).filter(ProofreadingSegmentBaseline.binding_id == binding.id).all()
        for baseline, segment in items:
            ordered_items.append((
                baseline.sheet_index,
                baseline.row_index,
                binding.target_column,
                int(segment.sequence_index or 0),
                baseline,
                segment,
            ))
    rows = []
    for order, (_, _, _, _, baseline, segment) in enumerate(sorted(ordered_items, key=lambda item: item[:4])):
        rows.append(ProofreadingExportRow(
            order=order,
            kind="缺译" if not segment.target_text else "对齐",
            source_text=segment.source_text,
            original_target_text=baseline.original_target_text,
            reviewed_target_text=segment.target_text or "",
            changed=(segment.target_text or "") != (baseline.original_target_text or ""),
            confirmation_status="已确认" if segment.status == "confirmed" else "未确认",
            llm_status="已校对" if segment.llm_provider else "未校对",
            confidence=None,
            method="xlsx",
            block_type=segment.block_type,
            block_index=segment.block_index,
            row_index=baseline.row_index,
            cell_index=segment.cell_index,
        ))
    return rows, file_record


def target_revision_export_available(batch: ProofreadingBatch) -> bool:
    """目标 DOCX 原件存在时，才允许生成保留目标排版的修订版。"""
    config = _batch_config(batch)
    target_filename = str(config.get("target_filename") or "")
    return (
        Path(target_filename).suffix.lower() == ".docx"
        and target_cache_path(batch.id, target_filename, create_parent=False).is_file()
    )


def build_export_readiness(db: Session, batch: ProofreadingBatch) -> dict[str, Any]:
    rows, file_record = build_proofreading_export_rows(db, batch)
    stage = getattr(batch, "workflow_stage", "not_applicable")
    if batch.batch_kind == "document_pair":
        available_formats = ["proofreading_audit_xlsx"]
        if stage == "proofreading":
            available_formats = [
                "proofreading_docx_layout",
                "proofreading_docx_ordered",
                "proofreading_audit_xlsx",
            ]
            if target_revision_export_available(batch):
                available_formats.insert(0, "proofreading_docx_target_revisions")
    else:
        available_formats = ["proofreading_xlsx_original"]
    return {
        "batch_id": str(batch.id),
        "total": len(rows),
        "confirmed": sum(row.confirmation_status == "已确认" for row in rows),
        "unconfirmed": sum(row.confirmation_status != "已确认" for row in rows),
        "missing_translation": sum(row.kind == "缺译" for row in rows),
        "translation_only": sum(row.kind == "增译" for row in rows),
        "translation_only_unreviewed": sum(row.kind == "增译" and row.llm_status == "未校对" for row in rows),
        "llm_failed": int(batch.failed_segments or 0),
        "available_formats": available_formats,
        "has_warnings": any((
            any(row.confirmation_status != "已确认" for row in rows),
            any(row.kind == "缺译" for row in rows),
            any(row.kind == "增译" and row.llm_status == "未校对" for row in rows),
            bool(batch.failed_segments),
        )),
    }


def _batch_config(batch: ProofreadingBatch) -> dict[str, Any]:
    try:
        value = json.loads(batch.config_json or "{}")
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _project_boundaries(before: str, after: str, boundaries: list[int]) -> list[int]:
    """把旧文本边界投影到新文本，保证边界单调且覆盖全部新文本。"""
    matcher = SequenceMatcher(a=before, b=after, autojunk=False)
    opcodes = matcher.get_opcodes()
    projected: list[int] = []
    for boundary in boundaries:
        mapped = len(after)
        for _tag, i1, i2, j1, j2 in opcodes:
            if boundary < i1:
                mapped = j1
                break
            if i1 <= boundary <= i2:
                if i2 == i1:
                    mapped = j2
                elif j2 == j1:
                    mapped = j1
                else:
                    ratio = (boundary - i1) / (i2 - i1)
                    mapped = j1 + round(ratio * (j2 - j1))
                break
        projected.append(max(0, min(len(after), mapped)))
    for index in range(1, len(projected)):
        projected[index] = max(projected[index], projected[index - 1])
    return projected


def _split_reviewed_text(original_parts: list[str], reviewed: str) -> list[str]:
    if not original_parts:
        return []
    if len(original_parts) == 1:
        return [reviewed]
    original = "".join(original_parts)
    boundaries: list[int] = []
    cursor = 0
    for part in original_parts[:-1]:
        cursor += len(part)
        boundaries.append(cursor)
    cuts = [0, *_project_boundaries(original, reviewed, boundaries), len(reviewed)]
    return [reviewed[cuts[index]:cuts[index + 1]] for index in range(len(original_parts))]


def _split_block_text(original_parts: list[str], joiner: str, reviewed: str) -> list[str]:
    """按原 DOCX 句段边界拆分块文本，不把解析时的人造连接符写回句段。"""
    if not original_parts:
        return []
    original = joiner.join(original_parts)
    spans: list[tuple[int, int]] = []
    cursor = 0
    for index, part in enumerate(original_parts):
        start = cursor
        end = start + len(part)
        spans.append((start, end))
        cursor = end + (len(joiner) if index < len(original_parts) - 1 else 0)
    boundary_values = [value for span in spans for value in span]
    mapped = _project_boundaries(original, reviewed, boundary_values)
    return [reviewed[mapped[index * 2]:mapped[index * 2 + 1]] for index in range(len(spans))]


def _workspace_block_key(segment: dict[str, Any]) -> tuple[int, int | None, int | None]:
    return (
        int(segment.get("block_index") or 0),
        segment.get("row_index"),
        segment.get("cell_index"),
    )


def _unit_block_key(unit: AlignUnit) -> tuple[int, int | None, int | None]:
    return (unit.block_index, unit.row_index, unit.cell_index)


def _build_target_revision_payload(
    db: Session,
    batch: ProofreadingBatch,
    target_bytes: bytes,
    target_filename: str,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """将对齐对的校对结果重新投影到目标 DOCX 的原始句段。"""
    from app.services.document_workspace import parse_docx_workspace

    config = _batch_config(batch)
    granularity = str(config.get("granularity") or "sentence")
    target_units = parse_side(target_bytes, target_filename, granularity)
    units_by_index = {unit.index: unit for unit in target_units}
    reviewed_by_pair = {
        row.pair_id: row.reviewed_target_text
        for row in build_proofreading_export_rows(db, batch)[0]
    }
    pairs = db.query(DocumentAlignmentPair).filter_by(batch_id=batch.id).order_by(
        DocumentAlignmentPair.pair_order,
    ).all()

    replacement_by_unit = {unit.index: unit.text for unit in target_units}
    deferred_insertions: list[tuple[int, str]] = []
    last_target_index: int | None = None
    for pair in pairs:
        try:
            indices = json.loads(pair.tgt_indices or "[]")
        except (TypeError, ValueError, json.JSONDecodeError):
            indices = []
        valid_indices: list[int] = []
        for index in indices if isinstance(indices, list) else []:
            try:
                normalized_index = int(index)
            except (TypeError, ValueError):
                continue
            if normalized_index in units_by_index:
                valid_indices.append(normalized_index)
        indices = valid_indices
        reviewed = str(reviewed_by_pair.get(str(pair.id), pair.target_text or "") or "")
        if indices:
            pieces = _split_reviewed_text([units_by_index[index].text for index in indices], reviewed)
            for index, piece in zip(indices, pieces):
                replacement_by_unit[index] = piece
            last_target_index = indices[-1]
        elif reviewed:
            # 原文有而目标文档无的段落没有天然版式锚点，先挂到前一目标单元；
            # 若位于文首，则在完成遍历后挂到后一目标单元之前。
            deferred_insertions.append((last_target_index if last_target_index is not None else -1, reviewed))

    for anchor_index, inserted in deferred_insertions:
        if anchor_index >= 0 and anchor_index in replacement_by_unit:
            replacement_by_unit[anchor_index] = f"{replacement_by_unit[anchor_index]}\n{inserted}"
        elif target_units:
            first = target_units[0].index
            replacement_by_unit[first] = f"{inserted}\n{replacement_by_unit[first]}"

    workspace_segments = list(parse_docx_workspace(target_bytes).get("segments", []))
    grouped: dict[tuple[int, int | None, int | None], list[dict[str, Any]]] = {}
    for segment in workspace_segments:
        grouped.setdefault(_workspace_block_key(segment), []).append(segment)
    units_by_block: dict[tuple[int, int | None, int | None], list[AlignUnit]] = {}
    for unit in target_units:
        units_by_block.setdefault(_unit_block_key(unit), []).append(unit)

    export_segments: list[dict[str, Any]] = []
    revisions: list[dict[str, Any]] = []
    for block_key, block_segments in grouped.items():
        block_type = str(block_segments[0].get("block_type") or "paragraph")
        if granularity == "sentence":
            source_parts = [
                str(item.get("display_text") or item.get("source_text") or "").strip()
                for item in block_segments
            ]
            joiner = "\n" if block_type == "table_cell" else " "
        else:
            source_parts = [
                str(item.get("display_text") or item.get("source_text") or "")
                for item in block_segments
            ]
            joiner = ""
        source_parts = [part for part in source_parts if part]
        old_block = joiner.join(source_parts).strip()
        edits = []
        for unit in units_by_block.get(block_key, []):
            edits.append((unit.source_start, unit.source_end, replacement_by_unit.get(unit.index, unit.text)))
        new_block = old_block
        for start, end, replacement in sorted(edits, reverse=True):
            new_block = f"{new_block[:start]}{replacement}{new_block[end:]}"

        corrected_parts = _split_block_text(source_parts, joiner, new_block)
        for segment, before_text, after_text in zip(block_segments, source_parts, corrected_parts):
            payload = dict(segment)
            payload["target_text"] = after_text
            payload["target_html"] = None
            export_segments.append(payload)
            sentence_id = str(segment.get("sentence_id") or "")
            if sentence_id and before_text != after_text:
                revisions.append({
                    "id": f"proofreading-{batch.id}-{sentence_id}",
                    "sentence_id": sentence_id,
                    "before_text": before_text,
                    "after_text": after_text,
                    "status": "pending",
                    "source": "manual",
                    "author": {"name": "校对工作流"},
                })
    return export_segments, revisions


def export_target_docx_with_revisions(
    db: Session, batch: ProofreadingBatch,
) -> tuple[bytes, str]:
    """以目标 DOCX 原件为母版，导出校对后的真实 Word 修订痕迹。"""
    from app.services.document_exporter import export_translated_docx

    if batch.batch_kind != "document_pair":
        raise ValueError("目标原格式修订导出仅用于双文档校对批次。")
    config = _batch_config(batch)
    target_filename = str(config.get("target_filename") or "")
    if Path(target_filename).suffix.lower() != ".docx":
        raise ValueError("目标原格式修订导出目前仅支持 DOCX 目标文档。")
    path = target_cache_path(batch.id, target_filename, create_parent=False)
    if not path.is_file():
        raise ValueError("该历史批次未保存目标文档原件，请重新导入双文档后再导出。")
    target_bytes = path.read_bytes()
    segments, revisions = _build_target_revision_payload(
        db, batch, target_bytes, target_filename,
    )
    content = export_translated_docx(
        target_bytes,
        segments,
        target_language=batch.target_language,
        revisions=revisions,
        include_revision_marks=True,
    )
    if not _docx_package_is_well_formed(content):
        raise ValueError("目标原格式修订文档生成失败，请检查目标 DOCX 后重试。")
    return content, f"{Path(target_filename).stem}_校对版_保留目标格式_含修订.docx"


def export_alignment_csv(pairs: list[DocumentAlignmentPair]) -> bytes:
    """导出对齐草稿；UTF-8 BOM 保证 Excel 直接打开中文不乱码。"""
    output = StringIO(newline="")
    fieldnames = [
        "pair_index", "source_indices", "target_indices", "source_text", "target_text",
        "confidence", "confidence_level", "method", "semantic_score", "status",
        "operation", "features_json",
    ]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for pair in pairs:
        features = json.loads(pair.features or "{}")
        if pair.source_text and not pair.target_text:
            status = "缺译"
        elif pair.target_text and not pair.source_text:
            status = "增译"
        elif pair.confidence_level == "low":
            status = "低置信"
        elif pair.confidence_level == "medium":
            status = "建议复核"
        else:
            status = "已对齐"
        writer.writerow({
            "pair_index": pair.pair_order,
            "source_indices": ",".join(map(str, json.loads(pair.src_indices or "[]"))),
            "target_indices": ",".join(map(str, json.loads(pair.tgt_indices or "[]"))),
            "source_text": pair.source_text,
            "target_text": pair.target_text,
            "confidence": pair.confidence,
            "confidence_level": pair.confidence_level,
            "method": pair.method,
            "semantic_score": features.get("semantic_similarity", features.get("absorbed_gap_similarity", "")),
            "status": status,
            "operation": features.get("op", ""),
            "features_json": json.dumps(features, ensure_ascii=False, separators=(",", ":")),
        })
    return ("\ufeff" + output.getvalue()).encode("utf-8")


def export_document_pair_xlsx(db: Session, batch: ProofreadingBatch) -> tuple[bytes, str]:
    """导出纯对齐结果，不读取或生成校对后译文。"""
    pairs = db.query(DocumentAlignmentPair).filter_by(batch_id=batch.id).order_by(
        DocumentAlignmentPair.pair_order,
    ).all()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "双文档对齐"
    sheet.append([
        "序号", "对齐状态", "原文", "译文", "置信等级", "置信度",
        "对齐方法", "原文单元索引", "译文单元索引",
    ])
    for cell in sheet[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(fill_type="solid", fgColor="FFD9EAF7")
    for pair in pairs:
        source_indices = json.loads(pair.src_indices or "[]")
        target_indices = json.loads(pair.tgt_indices or "[]")
        status = (
            "原文无对应译文" if source_indices and not target_indices
            else "译文无对应原文" if target_indices and not source_indices
            else "低置信" if pair.confidence_level == "low"
            else "建议复核" if pair.confidence_level == "medium"
            else "已对齐"
        )
        sheet.append([
            pair.pair_order + 1, status, pair.source_text or "", pair.target_text or "",
            pair.confidence_level, pair.confidence, pair.method,
            ",".join(map(str, source_indices)), ",".join(map(str, target_indices)),
        ])
        if status in {"原文无对应译文", "译文无对应原文", "低置信"}:
            cell = sheet.cell(row=sheet.max_row, column=2)
            cell.font = Font(color="FF9C0006", bold=True)
            cell.fill = PatternFill(fill_type="solid", fgColor="FFFFC7CE")
        for cell in sheet[sheet.max_row]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = f"A1:I{sheet.max_row}"
    sheet.sheet_view.showGridLines = False
    for column, width in zip("ABCDEFGHI", (9, 18, 48, 48, 12, 12, 20, 20, 20)):
        sheet.column_dimensions[column].width = width
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue(), f"{Path(batch.filename).stem}_原文译文对齐表.xlsx"


def export_ordered_bilingual_docx(
    db: Session, batch: ProofreadingBatch,
) -> tuple[bytes, str]:
    """生成顺序优先的原文在前双语 Word，不依赖源文件版式映射。"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    rows, _ = build_proofreading_export_rows(db, batch)
    document = Document()
    title = document.add_heading("双语校对文档", level=1)
    title.paragraph_format.keep_with_next = True
    for row in rows:
        source = document.add_paragraph()
        source.paragraph_format.keep_with_next = True
        source_run = source.add_run(row.source_text or TRANSLATION_ONLY_SOURCE_LABEL)
        source_run.bold = row.block_type == "heading"
        target = document.add_paragraph()
        target.paragraph_format.space_after = Pt(8)
        target_text = row.reviewed_target_text or MISSING_TRANSLATION_LABEL
        target_run = target.add_run(target_text)
        if row.kind == "增译":
            source_run.font.color.rgb = RGBColor(127, 96, 0)
        if row.kind == "缺译":
            target_run.bold = True
            target_run.font.color.rgb = RGBColor(156, 0, 6)
    output = BytesIO()
    document.save(output)
    return output.getvalue(), f"{Path(batch.filename).stem}_双语校对版_顺序优先.docx"


def _normalize_export_text(value: str) -> str:
    return "".join((value or "").split())


def _docx_package_is_well_formed(content: bytes) -> bool:
    """执行不依赖 Office 的基础 OOXML 包完整性检查。"""
    try:
        with ZipFile(BytesIO(content)) as archive:
            names = set(archive.namelist())
            if archive.testzip() is not None:
                return False
            if not {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}.issubset(names):
                return False
            for name in names:
                if name.endswith((".xml", ".rels")):
                    ET.fromstring(archive.read(name))
    except (BadZipFile, KeyError, ET.ParseError, OSError):
        return False
    return True


def _layout_export_is_complete(
    content: bytes,
    rows: list[ProofreadingExportRow],
    *,
    source_content: bytes | None = None,
    document_parse_mode: str = "full",
    document_parse_options: dict[str, object] | str | None = None,
) -> bool:
    """用目标文本顺序门禁拦截内容缺失或错序的保留排版结果。"""
    if not _docx_package_is_well_formed(content):
        return False
    try:
        from app.services.document_workspace import parse_docx_workspace

        parsed = parse_docx_workspace(
            content,
            document_parse_mode=document_parse_mode,
            document_parse_options=document_parse_options,
        )
        output_text = _normalize_export_text("\n".join(
            str(item.get("display_text") or item.get("source_text") or "")
            for item in parsed.get("segments", [])
        ))
    except Exception:
        return False
    cursor = 0
    expected_counts: dict[str, int] = {}
    for row in rows:
        expected = _normalize_export_text(row.reviewed_target_text or MISSING_TRANSLATION_LABEL)
        if not expected:
            continue
        expected_counts[expected] = expected_counts.get(expected, 0) + 1
        position = output_text.find(expected, cursor)
        if position < 0:
            return False
        cursor = position + len(expected)
    if not all(output_text.count(text) >= count for text, count in expected_counts.items()):
        return False
    if source_content is not None:
        try:
            source_workspace = parse_docx_workspace(
                source_content,
                document_parse_mode=document_parse_mode,
                document_parse_options=document_parse_options,
            )
        except Exception:
            return False
        source_output_text = compact_match_core(output_text)
        source_cursor = 0
        for segment in source_workspace.get("segments", []):
            expected_source = compact_match_core(str(
                segment.get("display_text") or segment.get("source_text") or ""
            ))
            if not expected_source:
                continue
            position = source_output_text.find(expected_source, source_cursor)
            if position < 0:
                return False
            source_cursor = position + len(expected_source)
    return True


def _build_layout_export_segments(
    db: Session,
    file_record: FileRecord,
    rows: list[ProofreadingExportRow],
) -> list[dict[str, Any]]:
    """把增译临时锚定到前一有效源文块，避免改写持久化的对齐位置。"""
    segment_by_pair: dict[str, Segment] = {}
    for segment in db.query(Segment).filter_by(file_record_id=file_record.id).all():
        pair_id = str(_segment_metadata(segment).get("alignment_pair_id") or "")
        if pair_id:
            segment_by_pair[pair_id] = segment

    ordered_segments = [segment_by_pair.get(row.pair_id) for row in rows]
    export_segments: list[dict[str, Any]] = []
    for index, (row, segment) in enumerate(zip(rows, ordered_segments)):
        if segment is None:
            continue
        metadata = _segment_metadata(segment)
        anchor = segment
        if bool(metadata.get("translation_only")):
            anchor = next(
                (
                    candidate for candidate in reversed(ordered_segments[:index])
                    if candidate is not None and not bool(_segment_metadata(candidate).get("translation_only"))
                ),
                None,
            ) or next(
                (
                    candidate for candidate in ordered_segments[index + 1:]
                    if candidate is not None and not bool(_segment_metadata(candidate).get("translation_only"))
                ),
                segment,
            )
        target_text = segment.target_text or MISSING_TRANSLATION_LABEL
        if bool(metadata.get("translation_only")):
            target_text = f"{TRANSLATION_ONLY_EXPORT_LABEL}{target_text}"
        export_segments.append({
            "sentence_id": segment.sentence_id,
            "source_text": segment.source_text,
            "display_text": segment.display_text,
            "target_text": target_text,
            "target_html": None,
            "source_html": segment.source_html,
            "numbering_text": str(metadata.get("numbering_text") or ""),
            "matched_source_text": segment.matched_source_text,
            "sequence_index": row.order,
            "block_type": anchor.block_type,
            "block_index": anchor.block_index,
            "row_index": anchor.row_index,
            "cell_index": anchor.cell_index,
            "segment_metadata": metadata,
        })
    return export_segments


def export_layout_bilingual_docx(
    db: Session, batch: ProofreadingBatch,
) -> tuple[bytes, str, bool]:
    """优先保留源 DOCX 排版；完整性门禁失败时返回顺序优先版本。"""
    from app.services.file_record_service import load_file_record_source
    from app.services.task_file_service import export_bilingual_task_docx_with_layout

    rows, file_record = build_proofreading_export_rows(db, batch)
    if file_record is None or Path(file_record.filename).suffix.lower() != ".docx":
        content, filename = export_ordered_bilingual_docx(db, batch)
        return content, filename, True
    raw_bytes = load_file_record_source(file_record)
    if raw_bytes is None:
        content, filename = export_ordered_bilingual_docx(db, batch)
        return content, filename, True
    export_segments = _build_layout_export_segments(db, file_record, rows)
    try:
        exported = export_bilingual_task_docx_with_layout(
            raw_bytes=raw_bytes,
            filename=file_record.filename,
            segments=export_segments,
            order="source_first",
            document_parse_mode=file_record.document_parse_mode,
            document_parse_options=file_record.document_parse_options,
            target_language=file_record.target_language,
        )
        if _layout_export_is_complete(
            exported.content,
            rows,
            source_content=raw_bytes,
            document_parse_mode=file_record.document_parse_mode,
            document_parse_options=file_record.document_parse_options,
        ):
            return (
                exported.content,
                f"{Path(batch.filename).stem}_双语校对版_保留排版.docx",
                False,
            )
    except Exception:
        pass
    content, filename = export_ordered_bilingual_docx(db, batch)
    return content, filename, True
