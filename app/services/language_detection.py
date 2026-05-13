from __future__ import annotations

import ast
import json
import re
from dataclasses import asdict, dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any

import yaml

from app.services.language_pairs import LANGUAGE_LABELS


MAX_TEXT_BYTES = 512 * 1024
MAX_SAMPLE_CHARS = 6000
MIN_SIGNAL_CHARS = 8

SUPPORTED_TEXT_EXTENSIONS = {
    ".csv",
    ".htm",
    ".html",
    ".json",
    ".md",
    ".markdown",
    ".po",
    ".pot",
    ".properties",
    ".srt",
    ".txt",
    ".yaml",
    ".yml",
}
SUPPORTED_DOCUMENT_EXTENSIONS = {".docx"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_DOCUMENT_EXTENSIONS

UNSUPPORTED_MESSAGE = "当前文件类型暂不支持自动识别，请手动选择源语言。"
NO_TEXT_MESSAGE = "未读取到足够文本，请手动选择源语言。"
UNCERTAIN_MESSAGE = "识别结果不确定，请手动选择源语言。"
SUCCESS_MESSAGE = "已识别源语言。"


@dataclass(frozen=True)
class LanguageDetectionResult:
    language: str | None
    label: str | None
    confidence: float
    supported: bool
    sample_length: int
    message: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth > 0:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self.parts.append(data)


def detect_upload_language(filename: str, raw_bytes: bytes) -> LanguageDetectionResult:
    extension = Path(filename or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return _empty_result(supported=False, message=UNSUPPORTED_MESSAGE)

    sample = _extract_sample_text(extension, raw_bytes)
    sample_length = len(re.sub(r"\s+", "", sample))
    if sample_length < MIN_SIGNAL_CHARS:
        return LanguageDetectionResult(
            language=None,
            label=None,
            confidence=0.0,
            supported=True,
            sample_length=sample_length,
            message=NO_TEXT_MESSAGE,
        )

    language, confidence = _detect_language(sample)
    if not language:
        return LanguageDetectionResult(
            language=None,
            label=None,
            confidence=0.0,
            supported=True,
            sample_length=sample_length,
            message=UNCERTAIN_MESSAGE,
        )

    return LanguageDetectionResult(
        language=language,
        label=LANGUAGE_LABELS.get(language),
        confidence=confidence,
        supported=True,
        sample_length=sample_length,
        message=SUCCESS_MESSAGE,
    )


def _empty_result(*, supported: bool, message: str) -> LanguageDetectionResult:
    return LanguageDetectionResult(
        language=None,
        label=None,
        confidence=0.0,
        supported=supported,
        sample_length=0,
        message=message,
    )


def _extract_sample_text(extension: str, raw_bytes: bytes) -> str:
    if extension == ".docx":
        return _extract_docx_text(raw_bytes)

    text = _decode_text(raw_bytes[:MAX_TEXT_BYTES])
    if not text:
        return ""

    if extension in {".html", ".htm"}:
        text = _extract_html_text(text)
    elif extension == ".json":
        text = _extract_json_text(text)
    elif extension in {".yaml", ".yml"}:
        text = _extract_yaml_text(text)
    elif extension == ".srt":
        text = _extract_srt_text(text)
    elif extension in {".po", ".pot"}:
        text = _extract_po_text(text)
    elif extension == ".properties":
        text = _extract_properties_text(text)
    elif extension in {".md", ".markdown"}:
        text = _extract_markdown_text(text)

    return _normalize_sample(text)


def _decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "gb18030", "big5", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def _extract_docx_text(raw_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(raw_bytes))
    except Exception:
        return ""

    parts: list[str] = []
    for paragraph in document.paragraphs:
        _append_text_part(parts, paragraph.text)
        if _joined_length(parts) >= MAX_SAMPLE_CHARS:
            return _normalize_sample(" ".join(parts))

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _append_text_part(parts, cell.text)
                if _joined_length(parts) >= MAX_SAMPLE_CHARS:
                    return _normalize_sample(" ".join(parts))

    return _normalize_sample(" ".join(parts))


def _append_text_part(parts: list[str], value: str | None) -> None:
    text = (value or "").strip()
    if text:
        parts.append(text)


def _joined_length(parts: list[str]) -> int:
    return sum(len(part) for part in parts)


def _extract_html_text(text: str) -> str:
    parser = _HTMLTextParser()
    try:
        parser.feed(text)
        return " ".join(parser.parts)
    except Exception:
        return re.sub(r"<[^>]+>", " ", text)


def _extract_json_text(text: str) -> str:
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return text

    parts: list[str] = []
    _collect_strings(data, parts)
    return " ".join(parts) if parts else text


def _extract_yaml_text(text: str) -> str:
    try:
        data = yaml.safe_load(text)
    except yaml.YAMLError:
        return text

    parts: list[str] = []
    _collect_strings(data, parts)
    return " ".join(parts) if parts else text


def _collect_strings(value: Any, parts: list[str]) -> None:
    if _joined_length(parts) >= MAX_SAMPLE_CHARS:
        return
    if isinstance(value, str):
        _append_text_part(parts, value)
        return
    if isinstance(value, dict):
        for item in value.values():
            _collect_strings(item, parts)
        return
    if isinstance(value, (list, tuple)):
        for item in value:
            _collect_strings(item, parts)


def _extract_srt_text(text: str) -> str:
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.isdigit() or "-->" in stripped:
            continue
        lines.append(stripped)
    return " ".join(lines)


def _extract_po_text(text: str) -> str:
    matches = re.findall(r'^(?:msgid|msgstr)\s+"(.*)"', text, flags=re.MULTILINE)
    if not matches:
        return text
    return " ".join(_unescape_po_string(match) for match in matches if match)


def _unescape_po_string(value: str) -> str:
    try:
        return str(ast.literal_eval(f'"{value}"'))
    except (SyntaxError, ValueError):
        return value


def _extract_properties_text(text: str) -> str:
    parts = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith(("#", "!")):
            continue
        if "=" in stripped:
            _, value = stripped.split("=", 1)
        elif ":" in stripped:
            _, value = stripped.split(":", 1)
        else:
            value = stripped
        _append_text_part(parts, value)
    return " ".join(parts)


def _extract_markdown_text(text: str) -> str:
    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
    text = re.sub(r"`[^`]+`", " ", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", " ", text, flags=re.MULTILINE)
    return text


def _normalize_sample(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", " ", text)
    text = re.sub(r"https?://\S+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:MAX_SAMPLE_CHARS]


def _detect_language(sample: str) -> tuple[str | None, float]:
    script_result = _detect_script_language(sample)
    if script_result[0]:
        return script_result
    return _detect_latin_language(sample)


def _detect_script_language(sample: str) -> tuple[str | None, float]:
    total_letters = sum(1 for char in sample if char.isalpha())
    if total_letters == 0:
        return None, 0.0

    kana_count = _count_chars_in_ranges(sample, (("\u3040", "\u30ff"),))
    if kana_count >= 3:
        return "ja-JP", _script_confidence(kana_count, total_letters)

    script_checks = (
        ("ko-KR", _count_chars_in_ranges(sample, (("\uac00", "\ud7af"),))),
        ("ru-RU", _count_chars_in_ranges(sample, (("\u0400", "\u04ff"),))),
        ("ar-SA", _count_chars_in_ranges(sample, (("\u0600", "\u06ff"),))),
        ("th-TH", _count_chars_in_ranges(sample, (("\u0e00", "\u0e7f"),))),
    )
    for language, count in script_checks:
        if count >= 3:
            return language, _script_confidence(count, total_letters)

    han_count = _count_chars_in_ranges(sample, (("\u4e00", "\u9fff"),))
    if han_count >= 4 and han_count / max(total_letters, 1) >= 0.18:
        return "zh-CN", _script_confidence(han_count, total_letters)

    return None, 0.0


def _count_chars_in_ranges(sample: str, ranges: tuple[tuple[str, str], ...]) -> int:
    return sum(1 for char in sample if any(start <= char <= end for start, end in ranges))


def _script_confidence(count: int, total_letters: int) -> float:
    ratio = count / max(total_letters, 1)
    return round(min(0.98, 0.66 + ratio * 0.3), 2)


LATIN_STOPWORDS: dict[str, set[str]] = {
    "en-US": {
        "a",
        "an",
        "and",
        "are",
        "as",
        "be",
        "by",
        "for",
        "from",
        "in",
        "is",
        "not",
        "of",
        "on",
        "or",
        "that",
        "the",
        "this",
        "to",
        "with",
    },
    "fr-FR": {
        "avec",
        "ce",
        "cette",
        "dans",
        "de",
        "des",
        "du",
        "est",
        "et",
        "la",
        "le",
        "les",
        "pas",
        "plus",
        "pour",
        "que",
        "qui",
        "sur",
        "un",
        "une",
    },
    "de-DE": {
        "auf",
        "das",
        "dem",
        "den",
        "der",
        "des",
        "die",
        "ein",
        "eine",
        "für",
        "im",
        "ist",
        "mit",
        "nicht",
        "und",
        "von",
        "zu",
    },
    "es-ES": {
        "como",
        "con",
        "de",
        "del",
        "el",
        "en",
        "es",
        "la",
        "las",
        "los",
        "para",
        "por",
        "que",
        "un",
        "una",
        "y",
    },
    "pt-BR": {
        "a",
        "as",
        "com",
        "da",
        "de",
        "do",
        "e",
        "em",
        "não",
        "o",
        "os",
        "para",
        "por",
        "que",
        "um",
        "uma",
    },
    "it-IT": {
        "che",
        "con",
        "del",
        "della",
        "di",
        "e",
        "gli",
        "il",
        "in",
        "la",
        "le",
        "lo",
        "non",
        "per",
        "un",
        "una",
    },
    "vi-VN": {
        "các",
        "cho",
        "của",
        "có",
        "không",
        "là",
        "một",
        "người",
        "như",
        "này",
        "trong",
        "và",
        "với",
        "được",
    },
}

DIACRITIC_HINTS: tuple[tuple[str, str], ...] = (
    ("vi-VN", r"[ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ]"),
    ("fr-FR", r"[àâæçéèêëîïôœùûüÿ]"),
    ("es-ES", r"[áéíóúñü¿¡]"),
    ("pt-BR", r"[ãõáâàéêíóôúç]"),
    ("de-DE", r"[äöüß]"),
    ("it-IT", r"[àèéìíîòóùú]"),
)


def _detect_latin_language(sample: str) -> tuple[str | None, float]:
    tokens = [token.lower() for token in re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]+", sample)]
    if len(tokens) < 4:
        return None, 0.0

    limited_tokens = tokens[:300]
    scores = {
        language: sum(1 for token in limited_tokens if token in stopwords)
        for language, stopwords in LATIN_STOPWORDS.items()
    }
    lowered_sample = sample.lower()
    for language, pattern in DIACRITIC_HINTS:
        if re.search(pattern, lowered_sample):
            scores[language] += 3

    ranked = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    best_language, best_score = ranked[0]
    second_score = ranked[1][1] if len(ranked) > 1 else 0

    if best_score <= 0:
        if len(limited_tokens) >= 8 and _is_basic_latin(sample):
            return "en-US", 0.52
        return None, 0.0

    if best_score < 2 and len(limited_tokens) > 12:
        return None, 0.0

    coverage = best_score / max(len(limited_tokens), 1)
    margin = (best_score - second_score) / max(best_score, 1)
    confidence = 0.5 + min(0.3, coverage * 2.4) + margin * 0.18
    if best_language == "en-US" and _is_basic_latin(sample):
        confidence += 0.03
    return best_language, round(min(0.93, confidence), 2)


def _is_basic_latin(sample: str) -> bool:
    letters = [char for char in sample if char.isalpha()]
    if not letters:
        return False
    basic = sum(1 for char in letters if "a" <= char.lower() <= "z")
    return basic / len(letters) >= 0.92
